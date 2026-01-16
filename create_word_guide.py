
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_guide():
    doc = Document()

    # Title
    title = doc.add_heading('🎙️ Ses Analiz Sistemi - Sunum Rehberi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Bu belge, projenizin sunumunda kullanabileceğiniz teknik detayları ve işleyiş yapısını özetler.')

    # 1. Dashboard
    doc.add_heading('🛠️ Modül ve Fonksiyon Haritası', level=1)
    
    doc.add_heading('1. Dashboard (Ana Ekran)', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Buton'
    hdr_cells[1].text = 'Modül'
    hdr_cells[2].text = 'Fonksiyon'
    hdr_cells[3].text = 'Açıklama'
    
    data = [
        ['KAYDI BAŞLAT / DURDUR', 'gui.py', 'toggle_recording', 'audio_recorder.py aracılığıyla mikrofonu açar/kapatır.'],
        ['SES DOSYASI YÜKLE', 'gui.py', 'process_audio_file', 'Mevcut bir .wav veya .mp3 dosyasını sisteme aktarır.']
    ]
    for b, m, f, d in data:
        row_cells = table.add_row().cells
        row_cells[0].text = b
        row_cells[1].text = m
        row_cells[2].text = f
        row_cells[3].text = d

    # 2. Analiz Raporu
    doc.add_heading('2. Analiz Raporu Sekmesi', level=2)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = 'Buton'
    hdr_cells[1].text = 'Modül'
    hdr_cells[2].text = 'Fonksiyon'
    hdr_cells[3].text = 'Açıklama'
    
    data2 = [
        ['GPT-4o İLE ANALİZ ET', 'gui.py', 'run_analysis', 'Metni OpenAI API\'sine göndererek derinlemesine analiz yapar.'],
        ['GEMINI İLE ANALİZ ET', 'gui.py', 'run_gemini_analysis', 'Metni Google Gemini API\'sine gönderir.'],
        ['RAPORU DIŞA AKTAR', 'gui.py', 'export_results', 'report_generator.py kullanarak PDF veya Word raporu oluşturur.'],
        ['Özetle / Kritik Noktalar', 'gui.py', '_send_quick_chat', 'AI\'ya hızlı komutlar gönderir.'],
        ['Yanıtı Seslendir (TTS)', 'gui.py', '_speak_last_response', 'OpenAI tts-1 modeliyle AI yanıtını seslendirir.'],
        ['AI\'ya Sor (Soru-Cevap)', 'gui.py', 'ask_ai_question', 'Konuşma metni üzerinde interaktif soru-cevap yapar.']
    ]
    for b, m, f, d in data2:
        row_cells = table2.add_row().cells
        row_cells[0].text = b
        row_cells[1].text = m
        row_cells[2].text = f
        row_cells[3].text = d

    # 3. Dil Koçu
    doc.add_heading('3. Dil Koçu (Language Coach) Sekmesi', level=2)
    table3 = doc.add_table(rows=1, cols=4)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'Buton'
    hdr_cells[1].text = 'Modül'
    hdr_cells[2].text = 'Fonksiyon'
    hdr_cells[3].text = 'Açıklama'
    
    data3 = [
        ['DİL ANALİZİ BAŞLAT', 'gui.py', 'run_language_analysis', 'Konuşmanızı gramer, kelime bilgisi ve telaffuz açısından inceler.'],
        ['DÜZELTMELERİ SESLEN.', 'gui.py', '_speak_language_resp.', 'AI\'nın dil önerilerini sesli olarak dinletir.'],
        ['PDF RAPOR AL', 'gui.py', 'save_coach_pdf', 'Dil gelişim raporunu PDF olarak kaydeder.']
    ]
    for b, m, f, d in data3:
        row_cells = table3.add_row().cells
        row_cells[0].text = b
        row_cells[1].text = m
        row_cells[2].text = f
        row_cells[3].text = d

    doc.add_page_break()

    # Offline/Online
    doc.add_heading('🌐 Çevrimdışı (Offline) ve Çevrimiçi (Online) İşlemler', level=1)
    doc.add_paragraph('Uygulamanız hibrit bir mimariye sahiptir; temel işlemler yerel donanımınızda, zeka gerektiren işlemler bulutta gerçekleşir.')
    
    doc.add_heading('🏠 Çevrimdışı (Yerel) İşlemler', level=2)
    offline = [
        'Ses Kayıt: Mikrofondan gelen veriler yerel olarak işlenir (sounddevice).',
        'Transkripsiyon (STT): OpenAI Whisper modeli bilgisayarınızda (CPU/GPU) yüklüdür.',
        'Görselleştirme: Kelime bulutu ve Duygu Grafikleri yerel olarak oluşturulur.',
        'Raporlama: PDF ve Word belgelerinin oluşturulması tamamen yereldir.',
        'Gürültü Engelleme: Ses iyileştirme işlemleri yerel kütüphanelerle yapılır.'
    ]
    for item in offline:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('☁️ Çevrimiçi (Bulut / LLM) İşlemler', level=2)
    online = [
        'Metin Analizi: GPT-4o ve Gemini Flash modelleri internet üzerinden API ile çalışır.',
        'AI Yanıt Seslendirme (TTS): Metni sese dönüştürme işlemi OpenAI sunucularında gerçekleşir.',
        'Soru-Cevap (Chat): Kullanıcının sorularını anlamak ve cevaplamak için LLM\'ler kullanılır.'
    ]
    for item in online:
        doc.add_paragraph(item, style='List Bullet')

    # Libraries
    doc.add_heading('📚 Kullanılan Temel Kütüphaneler', level=1)
    libs = [
        'CustomTkinter: Modern arayüz.',
        'OpenAI-Whisper: Yerel STT motoru.',
        'Google-GenerativeAI: Gemini entegrasyonu.',
        'OpenAI API: GPT-4o ve TTS desteği.',
        'Matplotlib & WordCloud: Görsel raporlama.',
        'FPDF & Python-Docx: Rapor çıktısı.',
        'SoundDevice & SciPy: Ses işleme.',
        'PyWinStyles: Windows efektleri.'
    ]
    for i, lib in enumerate(libs, 1):
        doc.add_paragraph(f'{i}. {lib}')

    # Tip
    doc.add_heading('💡 Sunum İpucu:', level=1)
    p = doc.add_paragraph()
    p.add_run('"Uygulamanın en güçlü yanı, Whisper ile yerel donanımda çalışan gizlilik dostu bir transkripsiyon süreci sunarken, analiz kısmında dünyanın en gelişmiş LLM servislerini aynı anda kullanabilmesidir."').italic = True

    output_path = r'c:\Users\kingm\Documents\GitHub\Speech-to-text\SUNUM_REHBERI.docx'
    doc.save(output_path)
    print(f"File saved to {output_path}")

if __name__ == "__main__":
    create_guide()
