"""
gui.py - Ana Kullanıcı Arayüzü (GUI) Modülü
Bu modül, uygulamanın görsel arayüzünü (CustomTkinter), ses kayıt kontrollerini, 
API entegrasyonlarını ve raporlama özelliklerini bir araya getirir.
"""

import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox, filedialog, ttk
import threading
import queue
import time
import sounddevice as sd
import soundfile as sf
import json
import torch
import os
import numpy as np
from openai import OpenAI
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from gemini_client import GeminiClient
from dotenv import load_dotenv, set_key
import datetime
import whisper
import noisereduce as nr
import pygame
import shutil
import pywinstyles # Modern Windows pencere efektleri için
from PIL import Image
import requests
import io

# .env dosyasını yükle (API anahtarları için)
load_dotenv()

# Karakter hatalarını önlemek için sistem dilini UTF-8 yapıyoruz
os.environ["PYTHONIOENCODING"] = "utf-8"

# Dinamik modül yüklemeleri (Opsiyonel bileşenler)
try:
    from analytics import AnalyticsGenerator
    from report_generator import ReportGenerator
    from visualizer import AudioVisualizer
    from elevenlabs_manager import ElevenLabsManager
    from stats_manager import StatsManager
    from sound_manager import SoundManager
except ImportError:
    # Eğer bu dosyalar mevcut değilse uygulama hatasız çalışmaya devam eder
    AnalyticsGenerator = None
    ReportGenerator = None
    AudioVisualizer = None
    ElevenLabsManager = None
    ElevenLabsManager = None
    from stats_manager import StatsManager # Always try to import local one
    from sound_manager import SoundManager

class SentimentTimeline(ctk.CTkFrame):
    """Analiz sekmesi için etkileşimli duygu zaman çizelgesi."""
    def __init__(self, master, textbox_to_scroll, **kwargs):
        super().__init__(master, **kwargs)
        self.textbox = textbox_to_scroll
        self.segments = []
        self.canvas = ctk.CTkCanvas(self, height=40, bg="#1a1a1a", highlightthickness=0)
        self.canvas.pack(fill="x", padx=10, pady=5)
        # Tıklama olayı geri yüklendi (User request)
        self.canvas.bind("<Button-1>", self._on_click)
        self.tooltip = None

    def update_timeline(self, segments):
        """
        segments: list of dicts like [{"text": "...", "sentiment": "pos/neg/neu", "index": float}]
        """
        self.segments = segments
        self.canvas.delete("all")
        if not segments: return

        width = self.canvas.winfo_width()
        if width <= 1: width = 600 # Fallback width

        total_length = sum(len(s["text"]) for s in segments)
        current_x = 0
        
        colors = {"pos": "#2ecc71", "neg": "#e74c3c", "neu": "#95a5a6"}
        
        for i, seg in enumerate(segments):
            seg_len = len(seg["text"])
            seg_width = (seg_len / total_length) * width
            
            x1 = current_x
            x2 = current_x + seg_width
            
            color = colors.get(seg.get("sentiment", "neu"), "#95a5a6")
            self.canvas.create_rectangle(x1, 5, x2, 35, fill=color, outline="", tags=f"seg_{i}")
            
            current_x += seg_width

    def _on_click(self, event):
        if not self.segments: return
        
        width = self.canvas.winfo_width()
        click_ratio = event.x / width
        
        total_text = "".join(s["text"] for s in self.segments)
        target_char_idx = int(click_ratio * len(total_text))
        
        # Metin kutusunda ilgili bölgeye ilerle
        current_char_count = 0
        for seg in self.segments:
            current_char_count += len(seg["text"])
            if current_char_count >= target_char_idx:
                # Metni bul ve yanıp sönme efektini yap (Opsiyonel)
                search_text = seg["text"][:30] # İlk 30 karakteri ara
                idx = self.textbox.search(search_text, "1.0", "end")
                if idx:
                    self.textbox.see(idx)
                    self.textbox.tag_add("highlight", idx, f"{idx} + {len(search_text)} chars")
                    self.textbox.tag_config("highlight", background="#00adb5", foreground="white")
                    self.after(1000, lambda: self.textbox.tag_remove("highlight", "1.0", "end"))
                break

class MicroAnimation:
    """Durum çubuğu için küçük, şık animasyonlar."""
    def __init__(self, parent_label):
        self.label = parent_label
        self.original_text = parent_label.cget("text")
        self.anim_running = False
        self.dots = 0

    def start_loading(self, text=None):
        if text: self.original_text = text
        self.anim_running = True
        self._animate_dots()

    def start_pulse(self):
        self.anim_running = True
        self._animate_pulse(0)

    def stop(self, final_text="Sistem Hazır"):
        self.anim_running = False
        self.label.configure(text=final_text)

    def _animate_dots(self):
        if not self.anim_running: return
        self.dots = (self.dots + 1) % 4
        self.label.configure(text=f"{self.original_text}{'.' * self.dots}")
        self.label.after(500, self._animate_dots)

    def _animate_pulse(self, step):
        if not self.anim_running: return
        # CustomTkinter'da direct alpha yok, alternatif renk değişimi
        colors = ["#00adb5", "#008a91", "#00676d", "#004449", "#00676d", "#008a91"]
        self.label.configure(text_color=colors[step % len(colors)])
        self.label.after(150, lambda: self._animate_pulse(step + 1))

class App(ctk.CTk):
    """
    Uygulamanın kalbi olan ana sınıfımız. Tüm pencere düzenini, ses kayıt süreçlerini,
    yapay zeka analizlerini ve raporlama sistemini bu sınıf üzerinden yönetiyoruz.
    """
    def __init__(self):
        super().__init__()
        
        # Donanım ve Durum Ayarları
        # Eğer NVIDIA GPU (CUDA) varsa kullan, yoksa CPU kullan
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.is_recording = False
        self.audio_frames = [] # Kayıt sırasında ses verilerinin toplandığı liste
        self.api_key = "" # OpenAI key
        self.fs = 16000 # Whisper için standart örnekleme hızı (Sample Rate)
        self.selected_mic_index = self.get_default_mic()
        self.audio_queue = queue.Queue() # Ses verileri için iş parçacığı güvenli kuyruk
        self.all_session_transcripts = [] # Oturum boyuncaki tüm transkriptleri saklayan liste
        self.recording_buttons = [] # Bu artık otomatik eşleme için kullanılmayacak, ama referans için kalsın
        self.active_recording_source = "home" # "home" veya "language"
        
        # Pygame Mixer Başlat (TTS ve Çalma için)
        try:
            pygame.mixer.init()
        except:
            print("Pygame mixer başlatılamadı.")
        
        # Son Analiz ve Transkript Verileri
        self.last_analysis = ""
        self.last_transcript = ""
        self.sentiment_stats = {'pos': 33, 'neg': 33, 'neu': 34}
        self.gemini_api_key = ""
        
        # Whisper Model Önbelleği
        self.whisper_model = None
        self.current_model_type = None
        
        # Dil Öğrenme (Language Coach) Durumu
        self.target_language = "İngilizce"
        self.user_level = "A2 (Gelişmekte Olan)"
        self.coach_mode = "Serbest Konuşma"
        self.language_analysis_result = ""
        self.coach_chat_history = [] # Soru-Cevap geçmişini saklamak için
        self.topic_chat_history = [] # Konu bazlı chat geçmişi
        
        # Eğitim Asistanı Quiz ve Gelişmiş Özellikler
        self.is_quiz_active = False
        self.current_quiz_questions = []
        self.current_quiz_index = 0
        self.quiz_score = 0
        self.topic_flashcards = []
        self.auto_tts_topic_var = tk.BooleanVar(value=True) # Varsayılan olarak açık
        self.last_topic_response = ""
        
        # Analiz Sonuçlarını Saklama (Çoklu PDF raporu için)
        self.analysis_results = {"OpenAI": "", "Gemini": ""}
        self.all_sentiment_stats = {"OpenAI": None, "Gemini": None}

        # ElevenLabs Ses Klonlama Yöneticisi
        self.eleven_api_key = os.getenv("ELEVENLABS_API_KEY", "").strip()
        if ElevenLabsManager:
            self.eleven_manager = ElevenLabsManager(api_key=self.eleven_api_key)
        else:
            self.eleven_manager = None
        self.eleven_voices = [] # [[name, id], ...]

        # İstatistik Yöneticisi
        # İstatistik Yöneticisi
        self.stats_manager = StatsManager()
        
        # Sound Manager Başlat
        self.sound_manager = SoundManager()

        # --- KARAKTER SES VE STİL EŞLEŞTİRMELERİ ---
        self.character_voices = {
            "Fatih Sultan Mehmet": "onyx",
            "M.K. Atatürk": "onyx",
            "Napolyon": "echo",
            "Jül Sezar": "echo",
            "Kanuni": "onyx",
            "Cengiz Han": "echo",
            "Albert Einstein": "fable",
            "Stephen Hawking": "fable",
            "Marie Curie": "shimmer",
            "Aziz Sancar": "alloy",
            "Sokrates": "fable",
            "Nietzsche": "echo"
        }
        
        self.character_styles = {
            "Fatih Sultan Mehmet": "Bir Osmanlı Sultanı gibi konuş. Kelimelerin otoriter, bilge ve hafif arkaik (Osmanlı Türkçesi esintili) olsun. Öğrenciye 'Lalam' veya 'Bilesin ki' gibi ifadelerle hitap et. İstanbul'un fatihi olduğunu hissettir.",
            "M.K. Atatürk": "M.K. Atatürk gibi konuş. Vizyoner, kararlı, modern ve teşvik edici bir üslup kullan. Hitabetin güçlü olsun. Milli mücadele ruhunu yansıt.",
            "Albert Einstein": "Biraz dağınık ama dahi bir profesör gibi konuş. Karmaşık konuları basit analojilerle anlat, merak uyandır. 'Hayal gücü bilgiden daha önemlidir' felsefesini yansıt.",
            "Marie Curie": "Tutkulu ve azimli bir bilim insanı olarak konuş. Laboratuvar deneyimlerinden, radyoaktiviteden ve bilimin zorluklarından bahset.",
            "Sokrates": "Sadece sorular sorarak karşındakinin gerçeği kendi kendine bulmasını sağlayan Sokratik yöntemi kullan. Bilgeliğini alçakgönüllülükle harmanla."
        }

        # --- EĞİTİM VE SENARYO VERİLERİ ---
        self.scenarios_data = {
            "Matematik": {
                "🧠 Problem Çözme Yarışması": [],
                "🎓 Profesör Modu": ["Cahit Arf", "Ali Nesin", "Pisagor", "Öklid"],
                "🌍 Gerçek Hayat Uygulamaları": ["Köprü Mühendisi", "Finans Analisti", "Kriptolog"],
                "🦉 Sokratik Öğretmen": []
            },
            "Fizik": {
                "🧪 Deney Simülasyonu": ["Kuantum Mekaniği", "Termodinamik", "Optik"],
                "🌌 Evrenin Sırları": ["Albert Einstein", "Stephen Hawking", "Richard Feynman", "Newton", "Marie Curie"],
                "🚀 Mühendislik Problemleri": ["Elon Musk (Rocket Scientist)", "Uçak Mühendisi"],
                "🤔 Kavramsal Tartışma": []
            },
            "Kimya": {
                "⚗️ Laboratuvar Kazaları": [],
                "🧬 Moleküler Keşif": ["Aziz Sancar", "Marie Curie", "Dmitri Mendeleev", "Rosalind Franklin"],
                "💥 Patlayıcı Deneyler": ["Alfred Nobel"],
                "🍳 Mutfak Kimyası (Eğlenceli)": []
            },
            "Biyoloji": {
                "🦠 Hastalık Dedektifi": [],
                "🧬 Genetik Mühendisi": ["CRISPR Uzmanı", "Darwin", "Mendel"],
                "🌿 Doğa Gözlemcisi": [],
                "🧠 Nörobilim Uzmanı": []
            },
            "Yapay Zeka": {
                "🤖 Gelecek Senaryoları": ["Ütopik", "Distopik", "Gerçekçi"],
                "🧠 Etik Tartışma": ["Trolley Problemi", "Bilinç Sorunsalı"],
                "💻 Teknik Mülakat": ["NLP Uzmanı", "CV Uzmanı", "LLM Mimarı"],
                "🔮 Teknoloji Kahini": []
            },
            "Kodlama": {
                "🛑 Code Review (Sert)": ["Huysuz Senior Dev", "Temiz Kod Takıntılısı", "Performans Canavarı"],
                "💼 Google Mülakatı": ["Algoritma Sorusu", "Sistem Tasarımı"],
                "🐞 Bug Avcısı": [],
                "👶 Bana 5 Yaşındayım Gibi Anlat": []
            },
            "Tarih": {
                "👑 Liderle Görüşme": ["Fatih Sultan Mehmet", "M.K. Atatürk", "Napolyon", "Jül Sezar", "Kanuni", "Cengiz Han"],
                "⏳ Zaman Yolcusu": ["İstanbul'un Fethi (1453)", "Fransız İhtilali (1789)", "Kurtuluş Savaşı", "Ay'a İniş"],
                "📜 Alternatif Tarih": ["Ya Hitler Kazansaydı?", "Ya Roma Çökmeseydi?"],
                "🏛️ Müze Rehberi": []
            },
            "Felsefe": {
                "⚖️ Münazara (Debate)": [],
                "🤔 Düşünce Deneyi": ["Mağara Alegorisi (Platon)", "Gemisi (Theseus)"],
                "🧠 Filozofla Sohbet": ["Sokrates", "Nietzsche", "Kant", "Aristoteles", "Mevlana"],
                "😈 Şeytanın Avukatı": []
            },
            "RPG Oyunu": {
                "🏰 Tarihsel Macera": ["İstanbul'un Fethi'nde Casus", "Kurtuluş Savaşı'nda Haberci", "Orta Çağ Krallığı"],
                "🚀 Uzay Kolonisi": ["Mars'ta Hayatta Kalma", "Yabancı Gezegen Keşfi"],
                "🕵️ Detektiflik Bürosu": ["Gizemli Cinayet", "Siber Suçlar"],
                "🧟 Kıyamet Sonrası": ["Zombi İstilası", "Nükleer Kış"]
            }
        }

        # Ana Pencere Konfigürasyonu
        self.title("Ses Analiz Sistemi")
        self.geometry("1300x950")
        ctk.set_appearance_mode("dark") # Koyu tema varsayılan
        self.protocol("WM_DELETE_WINDOW", self.on_app_closing)
        
        # Auto-VAD (Silence Detection) Ayarları
        self.silence_threshold = 0.01 # Sessizlik eşiği (RMS)
        self.silence_start_time = None
        self.auto_vad_enabled = False # Kullanıcının isteği üzerine varsayılan olarak KAPALI
        self.last_rms = 0

        # Arayüzü oluştur ve kayıtlı anahtarları yükle
        self.setup_ui()
        self.load_api_key()

        # Animasyon Yöneticisi
        self.animator = MicroAnimation(self.status_label)
        
        # Windows Modern Efektlerini Uygula (Glassmorphism)
        try:
            # Arka planı koyu ve pürüzsüz yap
            pywinstyles.apply_style(self, "mica")
            # Sol menüye hafif bir opaklık ver
            pywinstyles.set_opacity(self.navigation_frame, value=0.9)
        except Exception as pe:
            print(f"Pencere stili hatası: {pe}")
        
        # --- Başlangıç Temizliği ---
        self.cleanup_temp_files()

    def cleanup_temp_files(self):
        """Uygulama başladığında eski geçici dosyaları temizler."""
        print("[*] Geçici dosyalar temizleniyor...")
        try:
            current_dir = os.path.dirname(os.path.abspath(__file__))
            for item in os.listdir(current_dir):
                if item.startswith("temp_") and (item.endswith(".wav") or item.endswith(".mp3") or item.endswith(".png")):
                    file_path = os.path.join(current_dir, item)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                    except Exception as e:
                        print(f"Dosya silinemedi {item}: {e}")
        except Exception as e:
            print(f"Temizlik sırasında hata: {e}")

    def get_default_mic(self):
        """Sistemdeki varsayılan mikrofonun indeksini bulur."""
        try:
            devices = sd.query_devices()
            for i, d in enumerate(devices):
                if d['max_input_channels'] > 0:
                    return i
        except:
            return None
        return None

    def get_mic_list(self):
        """Kullanılabilir mikrofonların listesini döner."""
        try:
            devices = sd.query_devices()
            return [f"{i}: {d['name']}" for i, d in enumerate(devices) if d['max_input_channels'] > 0]
        except:
            return ["Mikrofon Bulunamadı"]

    def setup_ui(self):
        """
        Görsel arayüzü (Arayüz Panelleri, Butonlar, Tablolar vb.) burada inşa ediyoruz.
        Tasarım olarak modern 'Cyberpunk Glassmorphism' stilini hedefledik.
        """
        # Kayıtlar klasörü yoksa oluştur
        if not os.path.exists("recordings"):
            os.makedirs("recordings")
            
        # Grid (Izgara) sistemini yapılandır (Sidebar ve Main Container)
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # --- YAN NAVİGASYON PANELİ (Navigation Sidebar) ---
        self.navigation_frame = ctk.CTkFrame(self, corner_radius=0)
        self.navigation_frame.grid(row=0, column=0, sticky="nsew")
        self.navigation_frame.grid_rowconfigure(5, weight=1)

        self.navigation_frame_label = ctk.CTkLabel(self.navigation_frame, text=" SES ANALİZ\nSİSTEMİ", 
                                                 font=ctk.CTkFont(family="Inter", size=20, weight="bold"),
                                                 text_color="#ff007f") # Neon Pink vurgu
        self.navigation_frame_label.grid(row=0, column=0, padx=20, pady=20)

        # Navigasyon Butonları
        self.home_button = ctk.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Dashboard",
                                        fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                        anchor="w", command=self.home_button_event)
        self.home_button.grid(row=1, column=0, sticky="ew")

        self.analysis_button = ctk.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Analiz Raporu",
                                            fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                            anchor="w", command=self.analysis_button_event)
        self.analysis_button.grid(row=2, column=0, sticky="ew")

        self.history_button = ctk.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Geçmiş Kayıtlar",
                                          fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                          anchor="w", command=self.history_button_event)
        self.history_button.grid(row=3, column=0, sticky="ew")

        self.language_button = ctk.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Dil Koçu (AI)",
                                            fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                            anchor="w", command=self.language_button_event)
        self.language_button.grid(row=4, column=0, sticky="ew")

        self.settings_button = ctk.CTkButton(self.navigation_frame, corner_radius=0, height=40, border_spacing=10, text="Ayarlar",
                                            fg_color="transparent", text_color=("gray10", "gray90"), hover_color=("gray70", "gray30"),
                                            anchor="w", command=self.settings_button_event)
        self.settings_button.grid(row=5, column=0, sticky="ew")

        # Görünüm Menüsü (Sidebar Alt Kısmı)
        self.appearance_mode_menu = ctk.CTkOptionMenu(self.navigation_frame, values=["Dark", "Light", "System"],
                                                    command=self.change_appearance_mode_event)
        self.appearance_mode_menu.grid(row=6, column=0, padx=20, pady=20, sticky="s")

        # --- ANA İÇERİK PANELLERİ ---
        
        # 1. DASHBOARD PANELİ (Ana Kayıt Ekranı)
        self.home_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.home_frame.grid_columnconfigure(0, weight=3) # Transkript ve Kontroller
        self.home_frame.grid_columnconfigure(1, weight=1) # İstatistikler
        self.home_frame.grid_rowconfigure(2, weight=1)

        # 1. Ses Görselleştirici (En Üst)
        self.viz_container = ctk.CTkFrame(self.home_frame, corner_radius=15)
        self.viz_container.grid(row=0, column=0, padx=10, pady=10, sticky="ew")
        
        if AudioVisualizer:
            self.visualizer = AudioVisualizer(self.viz_container, mode="neon_bars") # Modern neon barlar
            self.visualizer.pack(fill="x", padx=2, pady=5)
        else:
            ctk.CTkLabel(self.viz_container, text="Görselleştirici Modülü Yüklenemedi").pack(pady=20)

        # 2. Durum Çubuğu (Barın Altında)
        self.status_bar = ctk.CTkFrame(self.home_frame, height=40, corner_radius=10)
        self.status_bar.grid(row=1, column=0, padx=20, pady=(0, 10), sticky="ew")
        
        self.status_label = ctk.CTkLabel(self.status_bar, text="Sistem Hazır", text_color="#ff007f", font=("Inter", 13, "bold"))
        self.status_label.pack(side="left", padx=20)

        color = "#ff007f" if self.device == "cuda" else "#ffea00"
        ctk.CTkLabel(self.status_bar, text=f"Donanım: {self.device.upper()}", text_color=color).pack(side="right", padx=20)

        # Transkript Alanı
        self.textbox = ctk.CTkTextbox(self.home_frame, font=("Inter", 15), corner_radius=15, border_width=2, border_color="#ff007f")
        self.textbox.grid(row=2, column=0, padx=20, pady=10, sticky="nsew")

        # Kontrol Butonları (Dashboard)
        self.dashboard_controls = ctk.CTkFrame(self.home_frame, fg_color="transparent")
        self.dashboard_controls.grid(row=3, column=0, padx=20, pady=(0, 20), sticky="ew")
        self.dashboard_controls.grid_columnconfigure((0, 1), weight=1)

        self.record_btn = ctk.CTkButton(self.dashboard_controls, text="KAYDI BAŞLAT", fg_color="green", font=("Arial", 14, "bold"),
                                       height=50, command=lambda: self.toggle_recording(source="home"))
        self.record_btn.grid(row=0, column=0, padx=(0, 5), sticky="ew")

        self.file_btn = ctk.CTkButton(self.dashboard_controls, text="SES DOSYASI YÜKLE", fg_color="#34495e", font=("Arial", 14, "bold"),
                                     height=50, command=self.process_audio_file)
        self.file_btn.grid(row=0, column=1, padx=(5, 0), sticky="ew")

        # 3. İstatistik Paneli (Sağ Taraf)
        self.stats_panel = ctk.CTkScrollableFrame(self.home_frame, label_text="Öğrenme İstatistikleri", corner_radius=15, border_width=2, border_color="#ff007f")
        self.stats_panel.grid(row=0, column=1, rowspan=4, padx=(10, 20), pady=10, sticky="nsew")
        
        # İstatistik Etiketleri
        self.stat_labels = {}
        stats_info = [
            ("Oturum Sayısı", "total_sessions", "📁"),
            ("Toplam Kelime", "total_words", "✍️"),
            ("Çalışma Süresi (dk)", "learning_time_minutes", "⏱️"),
            ("Tamamlanan Quiz", "total_quizzes", "📝"),
            ("Ortalamas Quiz Skoru", "average_quiz_score", "🎯")
        ]
        
        for name, key, icon in stats_info:
            frame = ctk.CTkFrame(self.stats_panel, fg_color="transparent")
            frame.pack(fill="x", pady=5)
            ctk.CTkLabel(frame, text=f"{icon} {name}:", font=("Inter", 12, "bold")).pack(side="left", padx=5)
            lbl = ctk.CTkLabel(frame, text="0", font=("Inter", 12), text_color="#00adb5")
            lbl.pack(side="right", padx=5)
            self.stat_labels[key] = lbl
        
        # Rozetler / Başarılar Alanı
        ctk.CTkLabel(self.stats_panel, text="🏆 Başarılar", font=("Inter", 14, "bold"), text_color="#ffea00").pack(pady=(20, 10))
        self.achievement_frame = ctk.CTkFrame(self.stats_panel, fg_color="transparent")
        self.achievement_frame.pack(fill="x")
        
        self.update_stats_ui()

        # 2. ANALİZ PANELİ (Detaylı AI Geri Bildirimi)
        self.analysis_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.analysis_frame.grid_columnconfigure(0, weight=3) # Metin alanı
        self.analysis_frame.grid_columnconfigure(1, weight=2) # Görsel alanı
        self.analysis_frame.grid_rowconfigure(1, weight=1)

        ctk.CTkLabel(self.analysis_frame, text="YAPAY ZEKA ANALİZ SONUÇLARI", font=("Arial", 22, "bold")).grid(row=0, column=0, columnspan=2, pady=20)

        # Analiz Metin Kutusu
        self.analysis_textbox = ctk.CTkTextbox(self.analysis_frame, font=("Consolas", 14), corner_radius=15)
        self.analysis_textbox.grid(row=1, column=0, padx=(20, 10), pady=10, sticky="nsew")

        # Görsel Alanı (Pie Chart & WordCloud)
        self.viz_frame = ctk.CTkScrollableFrame(self.analysis_frame, corner_radius=15, label_text="Görsel Bilgi Kartları")
        self.viz_frame.grid(row=1, column=1, padx=(10, 20), pady=10, sticky="nsew")

        self.sentiment_img_label = ctk.CTkLabel(self.viz_frame, text="Duygu Analizi Henüz Yapılmadı")
        self.sentiment_img_label.pack(pady=10)

        # Sentiment Timeline (Yeni)
        self.timeline_label = ctk.CTkLabel(self.viz_frame, text="Zaman Bazlı Duygu Dağılımı (Tıklanabilir):", font=("Arial", 11, "bold"))
        self.timeline_label.pack(pady=(10, 0))
        self.sentiment_timeline = SentimentTimeline(self.viz_frame, self.analysis_textbox, height=50, fg_color="transparent")
        self.sentiment_timeline.pack(fill="x", padx=5)

        self.wordcloud_img_label = ctk.CTkLabel(self.viz_frame, text="Kelime Bulutu Henüz Oluşturulmadı")
        self.wordcloud_img_label.pack(pady=10)

        self.chat_display = ctk.CTkTextbox(self.analysis_frame, height=250)
        self.chat_display.grid(row=3, column=0, columnspan=2, padx=20, pady=10, sticky="nsew")

        # Hızlı Aksiyon Butonları
        self.btn_row = ctk.CTkFrame(self.analysis_frame, fg_color="transparent")
        self.btn_row.grid(row=4, column=0, columnspan=2, padx=20, pady=(0, 10), sticky="ew")
        self.btn_row.grid_columnconfigure((0, 1, 2), weight=1)
        
        self.btn_summary = ctk.CTkButton(self.btn_row, text="📋 Özetle", width=100, command=lambda: self._send_quick_chat("Bu konuşmanın kısa ve etkili bir özetini çıkar."))
        self.btn_summary.grid(row=0, column=0, padx=5, sticky="ew")
        
        self.btn_points = ctk.CTkButton(self.btn_row, text="🎯 Kritik Noktalar", width=120, command=lambda: self._send_quick_chat("Bu konuşmadaki en önemli 3 kritik noktayı maddeler halinde yaz."))
        self.btn_points.grid(row=0, column=1, padx=5, sticky="ew")
        
        self.btn_tts = ctk.CTkButton(self.btn_row, text="🔊 Yanıtı Seslendir", width=130, command=self._speak_last_response, fg_color="#ff5722", hover_color="#e64a19")
        self.btn_tts.grid(row=0, column=2, padx=5, sticky="ew")

        # Analiz Başlatma Butonları (Analiz Sekmesi Üstü)
        self.analysis_actions = ctk.CTkFrame(self.analysis_frame, fg_color="transparent")
        self.analysis_actions.grid(row=2, column=0, columnspan=2, padx=20, pady=(10, 0), sticky="ew")
        self.analysis_actions.grid_columnconfigure((0, 1, 2), weight=1)

        self.analyze_btn = ctk.CTkButton(self.analysis_actions, text="GPT-4o İLE ANALİZ ET", fg_color="#10a37f", height=45, command=self.run_analysis)
        self.analyze_btn.grid(row=0, column=0, padx=5, sticky="ew")

        self.gemini_analyze_btn = ctk.CTkButton(self.analysis_actions, text="GEMINI İLE ANALİZ ET", fg_color="#4285f4", height=45, command=self.run_gemini_analysis)
        self.gemini_analyze_btn.grid(row=0, column=1, padx=5, sticky="ew")

        self.export_btn = ctk.CTkButton(self.analysis_actions, text="RAPORU DIŞA AKTAR", fg_color="#e67e22", height=45, command=self.export_results)
        self.export_btn.grid(row=0, column=2, padx=5, sticky="ew")

        # --- AI CHAT (SORU-CEVAP) BÖLÜMÜ ---
        self.chat_frame = ctk.CTkFrame(self.analysis_frame, corner_radius=15, border_width=1, border_color="#ff007f")
        self.chat_frame.grid(row=5, column=0, columnspan=2, padx=20, pady=(0, 20), sticky="ew")
        
        ctk.CTkLabel(self.chat_frame, text="AI'ya Sor:", font=("Arial", 12, "bold")).pack(side="left", padx=10, pady=10)
        self.chat_entry = ctk.CTkEntry(self.chat_frame, placeholder_text="Bu konuşmadan ne öğrenmek istersin?", height=35)
        self.chat_entry.pack(side="left", fill="x", expand=True, padx=10, pady=10)
        
        self.ask_btn = ctk.CTkButton(self.chat_frame, text="SOR", width=80, height=35, fg_color="#ff2e63", command=self.ask_ai_question)
        self.ask_btn.pack(side="right", padx=10, pady=10)
        
        # Soru kutusunda Enter tuşuna basınca soruyu gönder
        self.chat_entry.bind("<Return>", lambda e: self.ask_ai_question())

        # ttk.Treeview stilini güncelle (CustomTkinter ile uyum için)
        from tkinter import ttk
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("Treeview", background="#2b2b2b", foreground="white", fieldbackground="#2b2b2b", borderwidth=0)
        style.map("Treeview", background=[('selected', '#ff007f')])

        # 3. GEÇMİŞ PANELİ
        self.history_frame = ctk.CTkFrame(self, fg_color="transparent")
        ctk.CTkLabel(self.history_frame, text="KAYIT GEÇMİŞİ", font=("Arial", 22, "bold")).pack(pady=20)
        
        # Geçmiş Tablosu
        self.history_table = ttk.Treeview(self.history_frame, columns=("Tarih", "Model", "Özet", "İşlem"), show="headings")
        self.history_table.heading("Tarih", text="Tarih")
        self.history_table.heading("Model", text="Model")
        self.history_table.heading("Özet", text="Özet")
        self.history_table.heading("İşlem", text="İşlem")
        self.history_table.pack(fill="both", expand=True, padx=20, pady=20)
        self.history_table.bind("<Double-1>", self._on_history_click) # Çift tıklama ile oynat
        self.history_table.bind("<<TreeviewSelect>>", self._on_history_click) # Seçimle de tetiklenebilir

        self.refresh_history_btn = ctk.CTkButton(self.history_frame, text="GEÇMİŞİ YENİLE", command=self.update_history_list)
        self.refresh_history_btn.pack(pady=20)

        # 4. DİL KOÇU (AI LANGUAGE COACH) PANELİ
        self.language_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.language_frame.grid_columnconfigure(0, weight=1)
        self.language_frame.grid_columnconfigure(1, weight=1) # Sağ panel için
        self.language_frame.grid_rowconfigure(1, weight=1)

        ctk.CTkLabel(self.language_frame, text="AI DİL KOÇU & EĞİTİM MERKEZİ", font=("Inter", 22, "bold"), text_color="#ff007f").grid(row=0, column=0, columnspan=2, pady=(20, 10))

        # --- SOL PANEL: DİL KOÇLUĞU ---
        self.coach_left_panel = ctk.CTkFrame(self.language_frame, fg_color="transparent")
        self.coach_left_panel.grid(row=1, column=0, padx=(20, 10), pady=10, sticky="nsew")
        self.coach_left_panel.grid_columnconfigure(0, weight=1)
        self.coach_left_panel.grid_rowconfigure(1, weight=1)

        # Dil Ayarları Üst Bar (Sol Panel İçinde)
        self.lang_coach_settings = ctk.CTkFrame(self.coach_left_panel)
        self.lang_coach_settings.grid(row=0, column=0, pady=10, sticky="ew")
        
        ctk.CTkLabel(self.lang_coach_settings, text="Dil:").pack(side="left", padx=5, pady=5)
        self.coach_lang_combo = ctk.CTkComboBox(self.lang_coach_settings, values=["İngilizce", "Almanca", "Fransızca", "İspanyolca", "İtalyanca", "Rusça"], width=100)
        self.coach_lang_combo.set("İngilizce")
        self.coach_lang_combo.pack(side="left", padx=2)

        ctk.CTkLabel(self.lang_coach_settings, text="Seviye:").pack(side="left", padx=5)
        self.coach_level_combo = ctk.CTkComboBox(self.lang_coach_settings, values=["A1", "A2", "B1", "B2", "C1"], width=70)
        self.coach_level_combo.set("A2")
        self.coach_level_combo.pack(side="left", padx=2)

        ctk.CTkLabel(self.lang_coach_settings, text="Mod:").pack(side="left", padx=5)
        self.coach_mode_combo = ctk.CTkComboBox(self.lang_coach_settings, values=["Serbest", "Gramatik", "Kelime"], width=90)
        self.coach_mode_combo.set("Serbest")
        self.coach_mode_combo.pack(side="left", padx=2)

        # Dil Koçu Geri Bildirim Alanı
        self.language_textbox = ctk.CTkTextbox(self.coach_left_panel, font=("Inter", 14), corner_radius=15, border_width=2, border_color="#ff007f")
        self.language_textbox.grid(row=1, column=0, pady=10, sticky="nsew")
        self.language_textbox.insert("1.0", "--- AI DİL KOÇU HAZIR ---\n")

        # Aksiyon Butonları (Sol Panel Altı)
        self.coach_actions = ctk.CTkFrame(self.coach_left_panel, fg_color="transparent")
        self.coach_actions.grid(row=3, column=0, pady=10, sticky="ew")
        self.coach_actions.grid_columnconfigure((0, 1, 2), weight=1)

        self.coach_record_btn = ctk.CTkButton(self.coach_actions, text="KAYDI BAŞLAT", fg_color="green", font=("Inter", 12, "bold"),
                                             height=40, command=lambda: self.toggle_recording(source="language"))
        self.coach_record_btn.grid(row=0, column=0, padx=2, sticky="ew")

        self.run_coach_btn = ctk.CTkButton(self.coach_actions, text="🚀 ANALİZ", fg_color="#ff007f", font=("Inter", 12, "bold"),
                                          height=40, command=self.run_language_analysis)
        self.run_coach_btn.grid(row=0, column=1, padx=2, sticky="ew")

        self.coach_pdf_btn = ctk.CTkButton(self.coach_actions, text="📄 PDF", fg_color="#e67e22", font=("Inter", 12, "bold"),
                                          height=40, command=self.save_coach_pdf)
        self.coach_pdf_btn.grid(row=0, column=2, padx=2, sticky="ew")

        self.word_bank_btn = ctk.CTkButton(self.coach_actions, text="📔 KELİME DEFTERİ", fg_color="#9b59b6", font=("Inter", 12, "bold"),
                                          height=40, command=self.show_word_bank)
        self.word_bank_btn.grid(row=1, column=0, pady=(5, 0), sticky="ew")

        self.pronounce_test_btn = ctk.CTkButton(self.coach_actions, text="🎯 TELAFFUZ TESTİ", fg_color="#2ecc71", font=("Inter", 12, "bold"),
                                               height=40, command=self.start_pronunciation_test)
        self.pronounce_test_btn.grid(row=1, column=1, pady=(5, 0), sticky="ew")

        self.speak_coach_btn = ctk.CTkButton(self.coach_actions, text="🔊 SESLENDİR", fg_color="#ff5722", font=("Inter", 12, "bold"),
                                            height=40, command=self._speak_language_response)
        self.speak_coach_btn.grid(row=1, column=2, pady=(5, 0), sticky="ew")

        # --- SAĞ PANEL: KONU BAZLI AI SOHBET ---
        self.topic_right_panel = ctk.CTkFrame(self.language_frame, fg_color="transparent")
        self.topic_right_panel.grid(row=1, column=1, padx=(10, 20), pady=10, sticky="nsew")
        self.topic_right_panel.grid_columnconfigure(0, weight=1)
        self.topic_right_panel.grid_rowconfigure(1, weight=1)

        # Konu Seçimi Üst Bar
        self.topic_settings = ctk.CTkFrame(self.topic_right_panel)
        self.topic_settings.grid(row=0, column=0, pady=10, sticky="ew")
        
        ctk.CTkLabel(self.topic_settings, text="Konu:").pack(side="left", padx=(10, 2), pady=10)
        self.topic_combo = ctk.CTkComboBox(self.topic_settings, values=list(self.scenarios_data.keys()), width=110, command=self._on_topic_change)
        self.topic_combo.set("Kodlama")
        self.topic_combo.pack(side="left", padx=2)

        # Yeni Senaryo ve Alt Seçim Kutuları
        self.scenario_combo = ctk.CTkComboBox(self.topic_settings, values=[], width=140, command=self._on_scenario_change)
        self.scenario_combo.set("Senaryo Seçiniz")
        self.scenario_combo.pack(side="left", padx=2)
        
        self.sub_option_combo = ctk.CTkComboBox(self.topic_settings, values=[], width=130)
        self.sub_option_combo.set("Karakter Seçiniz")
        # Başlangıçta gizli olabilir ama grid kullandığımız için pack_forget yapabiliriz, 
        # şimdilik varsayılan olarak gösterip boş bırakalım veya kodla yönetelim.
        self.sub_option_combo.pack(side="left", padx=2)

        self.magic_wand_btn = ctk.CTkButton(self.topic_settings, text="🪄", fg_color="#ff007f", font=("Inter", 16),
                                           command=self.create_custom_scenario, width=40)
        self.magic_wand_btn.pack(side="left", padx=5)

        # İlk Başlatma: Kodlama için senaryoları yükle
        self._on_topic_change("Kodlama")

        self.start_topic_btn = ctk.CTkButton(self.topic_settings, text="BAŞLAT", fg_color="#4285f4", font=("Inter", 12, "bold"),
                                            command=self.run_topic_ai_chat, width=70)
        self.start_topic_btn.pack(side="left", padx=5)

        # Konu Sohbet Alanı
        self.topic_textbox = ctk.CTkTextbox(self.topic_right_panel, font=("Consolas", 14), corner_radius=15, border_width=2, border_color="#4285f4")
        self.topic_textbox.grid(row=1, column=0, pady=5, sticky="nsew")
        
        # Görsel Alanı (DALL-E) - Metin kutusunun altında (Simetri için en iyisi bu)
        self.topic_image_frame = ctk.CTkFrame(self.topic_right_panel, height=350, fg_color="transparent")
        self.topic_image_frame.grid(row=2, column=0, pady=5, sticky="ew")
        
        self.topic_image_label = ctk.CTkLabel(self.topic_image_frame, text="", text_color="gray50")
        self.topic_image_label.pack(expand=True, fill="both")

        self.start_quiz_btn = ctk.CTkButton(self.topic_settings, text="📝 QUIZ", fg_color="#10a37f", font=("Inter", 12, "bold"),
                                           command=self.run_topic_quiz, width=50)
        self.start_quiz_btn.pack(side="left", padx=5)

        self.flashcard_btn = ctk.CTkButton(self.topic_settings, text="🎴 KARTLAR", fg_color="#ffea00", text_color="black", font=("Inter", 12, "bold"),
                                          command=self.generate_flashcards, width=70)
        self.flashcard_btn.pack(side="left", padx=5)

        self.topic_pdf_btn = ctk.CTkButton(self.topic_settings, text="📄 PDF", fg_color="#e67e22", font=("Inter", 12, "bold"),
                                          command=self.save_topic_pdf, width=50)
        self.topic_pdf_btn.pack(side="left", padx=5)

        self.image_gen_btn = ctk.CTkButton(self.topic_settings, text="🖼️", fg_color="#8e44ad", font=("Inter", 16),
                                          command=self.manual_image_generation, width=40)
        self.image_gen_btn.pack(side="left", padx=5)

        self.topic_rag_btn = ctk.CTkButton(self.topic_settings, text="📂 YÜKLE", fg_color="#3498db", font=("Inter", 12, "bold"),
                                          command=self.upload_topic_notes, width=70)
        self.topic_rag_btn.pack(side="left", padx=5)


        # RPG Envanter / Durum Paneli
        self.inv_frame = ctk.CTkFrame(self.topic_right_panel, fg_color="transparent")
        self.inv_frame.grid(row=2, column=0, pady=2, sticky="ne")
        
        # Envanter Label'ları
        self.hp_label = ctk.CTkLabel(self.inv_frame, text="", text_color="#e74c3c", font=("Impact", 18))
        self.hp_label.pack(side="top", anchor="e")
        
        self.inv_label = ctk.CTkLabel(self.inv_frame, text="", text_color="#f1c40f", font=("Inter", 12))
        self.inv_label.pack(side="top", anchor="e")

        self.topic_textbox.insert("1.0", "--- KONU BAZLI EĞİTİM ASİSTANI ---\nLütfen bir konu seçip 'SOHBETİ BAŞLAT' butonuna basın.\n")

        # Konu Chat Giriş Alanı
        self.topic_chat_input_frame = ctk.CTkFrame(self.topic_right_panel, corner_radius=15, border_width=1, border_color="#4285f4")
        self.topic_chat_input_frame.grid(row=3, column=0, pady=10, sticky="ew")
        
        self.topic_mic_btn = ctk.CTkButton(self.topic_chat_input_frame, text="🎤", width=35, height=35, fg_color="transparent", 
                                           text_color="#4285f4", font=("Arial", 16), command=lambda: self.toggle_recording(source="topic_chat"))
        self.topic_mic_btn.pack(side="left", padx=(10, 0), pady=10)

        self.topic_chat_entry = ctk.CTkEntry(self.topic_chat_input_frame, placeholder_text="Seçili konu hakkında bir şey sor...", height=35)
        self.topic_chat_entry.pack(side="left", fill="x", expand=True, padx=10, pady=10)
        
        self.quiz_option_frame = ctk.CTkFrame(self.topic_chat_input_frame, fg_color="transparent")
        self.quiz_options = {}
        for opt in ["A", "B", "C", "D"]:
            btn = ctk.CTkButton(self.quiz_option_frame, text=opt, width=40, font=("Inter", 12, "bold"),
                                command=lambda o=opt: self.submit_quiz_answer(o))
            btn.pack(side="left", padx=2)
            self.quiz_options[opt] = btn
            
        # RPG Seçenekleri Frame'i
        self.rpg_option_frame = ctk.CTkFrame(self.topic_chat_input_frame, fg_color="transparent")
        self.rpg_buttons = []
        
        self.topic_speak_btn = ctk.CTkButton(self.topic_chat_input_frame, text="🔊", width=35, height=35, fg_color="transparent",
                                             text_color="#10a37f", font=("Arial", 16), command=self._speak_topic_last_response)
        # Hoparlör butonunu GÖNDER'in yanına (soluna) ekleyelim
        self.topic_speak_btn.pack(side="right", padx=(0, 5), pady=10)

        # OTOTTS Anahtarı (GÖNDER'in yanına eklendi)
        self.auto_tts_topic_switch = ctk.CTkSwitch(self.topic_chat_input_frame, text="OTOTTS", variable=self.auto_tts_topic_var, 
                                                 font=("Inter", 11, "bold"), width=80)
        self.auto_tts_topic_switch.pack(side="right", padx=(5, 5), pady=10)

        self.topic_ask_btn = ctk.CTkButton(self.topic_chat_input_frame, text="GÖNDER", width=80, height=35, fg_color="#4285f4", command=self.run_topic_ai_chat)
        self.topic_ask_btn.pack(side="right", padx=(5, 10), pady=10)
        
        self.topic_chat_entry.bind("<Return>", lambda e: self.run_topic_ai_chat())

        # 5. AYARLAR PANELİ (Kaydırılabilir)
        self.settings_frame = ctk.CTkScrollableFrame(self, fg_color="transparent")
        
        ctk.CTkLabel(self.settings_frame, text="SİSTEM AYARLARI", font=("Arial", 22, "bold")).pack(pady=20)

        # API Ayarları Grubu
        self.api_group = ctk.CTkFrame(self.settings_frame)
        self.api_group.pack(padx=40, pady=10, fill="x")
        
        ctk.CTkLabel(self.api_group, text="API YAPILANDIRMASI", font=("Arial", 14, "bold")).pack(pady=10)
        
        ctk.CTkLabel(self.api_group, text="OpenAI API Anahtarı:").pack()
        self.api_entry = ctk.CTkEntry(self.api_group, width=400, show="*")
        self.api_entry.pack(pady=5)

        ctk.CTkLabel(self.api_group, text="Gemini API Anahtarı:").pack()
        self.gemini_api_entry = ctk.CTkEntry(self.api_group, width=400, show="*")
        self.gemini_api_entry.pack(pady=5)

        ctk.CTkButton(self.api_group, text="Anahtarları Güvenli Kaydet", command=self.save_api_keys).pack(pady=15)

        # Model Ayarları Grubu
        self.model_group = ctk.CTkFrame(self.settings_frame)
        self.model_group.pack(padx=40, pady=10, fill="x")
        
        ctk.CTkLabel(self.model_group, text="MODEL VE DONANIM", font=("Arial", 14, "bold")).pack(pady=10)

        # Grid for model settings
        model_grid = ctk.CTkFrame(self.model_group, fg_color="transparent")
        model_grid.pack(pady=5)

        ctk.CTkLabel(model_grid, text="Whisper Modeli:").grid(row=0, column=0, padx=10)
        self.model_combo = ctk.CTkComboBox(model_grid, values=["tiny", "base", "small", "medium", "large-v3"])
        self.model_combo.set("medium")
        self.model_combo.grid(row=0, column=1, pady=5)

        ctk.CTkLabel(model_grid, text="Kaynak Dil:").grid(row=1, column=0, padx=10)
        self.lang_options = {
            "Otomatik Algıla": None,
            "Türkçe": "turkish", 
            "İngilizce": "english", 
            "Almanca": "german", 
            "Fransızca": "french", 
            "İspanyolca": "spanish", 
            "İtalyanca": "italian", 
            "Rusça": "russian"
        }
        self.lang_combo = ctk.CTkComboBox(model_grid, values=list(self.lang_options.keys()))
        self.lang_combo.set("Türkçe")
        self.lang_combo.grid(row=1, column=1, pady=5)

        ctk.CTkLabel(model_grid, text="Mikrofon:").grid(row=2, column=0, padx=10)
        self.mic_combo = ctk.CTkComboBox(model_grid, values=self.get_mic_list(), command=self.change_mic)
        self.mic_combo.grid(row=2, column=1, pady=5)

        ctk.CTkLabel(model_grid, text="Yapay Zeka Sesi:").grid(row=3, column=0, padx=10)
        self.tts_voices = {
            "Profesyonel Erkek (Onyx)": "onyx",
            "Sert Erkek (Echo)": "echo",
            "Genç ve Nazik (Nova)": "nova",
            "Net ve Parlak (Shimmer)": "shimmer",
            "Dışavurumcu (Fable)": "fable",
            "Dengeli ve Nötr (Alloy)": "alloy"
        }
        self.tts_voice_combo = ctk.CTkComboBox(model_grid, values=list(self.tts_voices.keys()))
        self.tts_voice_combo.set("Genç ve Nazik (Nova)")
        self.tts_voice_combo.grid(row=3, column=1, pady=5)

        ctk.CTkLabel(model_grid, text="AI Karakteri:").grid(row=4, column=0, padx=10)
        self.personas = {
            "Profesyonel Analist": "analyst",
            "Sert Mentor": "strict_mentor",
            "Samimi Teknoloji Gurusu": "tech_guru",
            "Akademik Gözlemci": "scholar",
            "Utangaç ve Cıvıl Cıvıl": "shy_girl"
        }
        self.persona_combo = ctk.CTkComboBox(model_grid, values=list(self.personas.keys()))
        self.persona_combo.set("Profesyonel Analist")
        self.persona_combo.grid(row=4, column=1, pady=5)

        # Switchler
        self.translate_var = ctk.BooleanVar(value=False)
        ctk.CTkSwitch(self.model_group, text="Tanımadan Sonra İngilizceye Çevir", variable=self.translate_var).pack(pady=5)
        
        self.autosave_var = ctk.BooleanVar(value=True)
        ctk.CTkSwitch(self.model_group, text="Ses Kayıtlarını Otomatik Arşivle", variable=self.autosave_var).pack(pady=5)

        self.noise_reduce_var = ctk.BooleanVar(value=True)
        ctk.CTkSwitch(self.model_group, text="Gelişmiş Gürültü Azaltma (Önerilen)", variable=self.noise_reduce_var).pack(pady=5)

        self.auto_vad_var = ctk.BooleanVar(value=False)
        ctk.CTkSwitch(self.model_group, text="Otomatik Sessizlik Algılama (Auto-VAD)", variable=self.auto_vad_var, command=self._toggle_auto_vad).pack(pady=5)

        ctk.CTkLabel(self.model_group, text="VAD Hassasiyeti (Daha düşük = Daha hassas):", font=("Arial", 11)).pack(pady=(5, 0))
        self.vad_threshold_slider = ctk.CTkSlider(self.model_group, from_=0.001, to=0.1, number_of_steps=100, command=self._update_vad_threshold)
        self.vad_threshold_slider.set(self.silence_threshold)
        self.vad_threshold_slider.pack(pady=5, padx=20)

        # ElevenLabs Ses Klonlama Grubu
        self.eleven_group = ctk.CTkFrame(self.settings_frame)
        self.eleven_group.pack(padx=40, pady=10, fill="x")
        
        ctk.CTkLabel(self.eleven_group, text="ELEVENLABS SES KLONLAMA (TTS)", font=("Arial", 14, "bold")).pack(pady=10)
        
        eleven_grid = ctk.CTkFrame(self.eleven_group, fg_color="transparent")
        eleven_grid.pack(pady=5)
        
        ctk.CTkLabel(eleven_grid, text="API Anahtarı:").grid(row=0, column=0, padx=10)
        self.eleven_api_entry = ctk.CTkEntry(eleven_grid, width=300, show="*")
        self.eleven_api_entry.insert(0, self.eleven_api_key)
        self.eleven_api_entry.grid(row=0, column=1, pady=5)
        
        self.eleven_enable_var = ctk.BooleanVar(value=False)
        self.eleven_switch = ctk.CTkSwitch(eleven_grid, text="ElevenLabs TTS Kullan", variable=self.eleven_enable_var)
        self.eleven_switch.grid(row=1, column=0, columnspan=2, pady=10)
        
        ctk.CTkLabel(eleven_grid, text="Seçili Ses:").grid(row=2, column=0, padx=10)
        self.eleven_voice_combo = ctk.CTkComboBox(eleven_grid, values=["Sesler Yükleniyor..."], width=200)
        self.eleven_voice_combo.grid(row=2, column=1, pady=5)
        
        ctk.CTkButton(self.eleven_group, text="Ses Listesini Güncelle / Bağlan", command=self._refresh_eleven_voices).pack(pady=10)
        
        # İlk yükleme
        self.after(2000, self._refresh_eleven_voices)

        # Varsayılan Sayfayı Göster
        self.select_frame_by_name("home")

    def select_frame_by_name(self, name):
        # Buton renklerini sıfırla
        self.home_button.configure(fg_color=("gray75", "gray25") if name == "home" else "transparent")
        self.analysis_button.configure(fg_color=("gray75", "gray25") if name == "analysis" else "transparent")
        self.history_button.configure(fg_color=("gray75", "gray25") if name == "history" else "transparent")
        self.language_button.configure(fg_color=("gray75", "gray25") if name == "language" else "transparent")
        self.settings_button.configure(fg_color=("gray75", "gray25") if name == "settings" else "transparent")

        # Sayfaları gizle
        self.home_frame.grid_forget()
        self.analysis_frame.grid_forget()
        self.history_frame.grid_forget()
        self.language_frame.grid_forget()
        self.settings_frame.grid_forget()

        # Seçilen sayfayı göster
        if name == "home":
            self.home_frame.grid(row=0, column=1, sticky="nsew")
        elif name == "analysis":
            self.analysis_frame.grid(row=0, column=1, sticky="nsew")
        elif name == "history":
            self.history_frame.grid(row=0, column=1, sticky="nsew")
            self.update_history_list()
        elif name == "language":
            self.language_frame.grid(row=0, column=1, sticky="nsew")
        elif name == "settings":
            self.settings_frame.grid(row=0, column=1, sticky="nsew")

    def home_button_event(self):
        self.select_frame_by_name("home")

    def analysis_button_event(self):
        self.select_frame_by_name("analysis")

    def history_button_event(self):
        self.select_frame_by_name("history")

    def language_button_event(self):
        self.select_frame_by_name("language")

    def settings_button_event(self):
        self.select_frame_by_name("settings")

    def change_appearance_mode_event(self, new_appearance_mode):
        ctk.set_appearance_mode(new_appearance_mode)

    def update_history_list(self):
        """Kayıtlar klasöründeki dosyaları listeler ve tabloyu günceller."""
        # Tabloyu temizle
        for item in self.history_table.get_children():
            self.history_table.delete(item)
            
        if not os.path.exists("recordings"):
            os.makedirs("recordings")
            
        recordings = sorted(os.listdir("recordings"), reverse=True)
        recordings = [f for f in recordings if f.endswith(".wav")]
        
        for filename in recordings:
            # Tarih bilgisini dosyadan çıkar (Format: kayit_20231227_120000.wav)
            date_info = filename.replace("kayit_", "").replace(".wav", "").replace("_", " ")
            self.history_table.insert("", "end", values=(
                date_info, 
                "Whisper", 
                f"{filename}", 
                "OYNAT ▶️"
            ))

    def update_stats_ui(self):
        """İstatistikleri arayüzde günceller."""
        if not hasattr(self, 'stats_manager') or not hasattr(self, 'stat_labels'):
            return
            
        stats = self.stats_manager.get_summary()
        for key, lbl in self.stat_labels.items():
            val = stats.get(key, 0)
            if key == "average_quiz_score":
                lbl.configure(text=f"%{val}")
            else:
                lbl.configure(text=str(val))
        
        # Başarıları (Achievements) güncelle
        for widget in self.achievement_frame.winfo_children():
            widget.destroy()
            
        if stats["total_sessions"] >= 5:
            ctk.CTkLabel(self.achievement_frame, text="🥉 Bronz Öğrenci", text_color="#cd7f32").pack()
        if stats["total_quizzes"] >= 1:
            ctk.CTkLabel(self.achievement_frame, text="🎓 Bilgi Avcısı", text_color="#00adb5").pack()

    def _toggle_auto_vad(self):
        """Auto-VAD özelliğini açar/kapatır."""
        self.auto_vad_enabled = self.auto_vad_var.get()
        status = "Açık" if self.auto_vad_enabled else "Kapalı"
        print(f"Auto-VAD: {status}")

    def _update_vad_threshold(self, value):
        """VAD hassasiyetini günceller."""
        self.silence_threshold = float(value)
        # print(f"VAD Eşiği Güncellendi: {self.silence_threshold}")


    def _on_history_click(self, event):
        """Geçmiş tablosuna tıklandığında kaydı oynatır."""
        selected = self.history_table.selection()
        if not selected: return
        
        item_values = self.history_table.item(selected[0])["values"]
        filename = item_values[2] # Özet/Filename kolonu
        path = os.path.join("recordings", filename)
        
        if os.path.exists(path):
            self._play_audio(path)
        else:
            messagebox.showinfo("Bilgi", "Ses dosyası bulunamadı.")

    def _play_audio(self, file_path):
        """Verilen ses dosyasını pygame ile çalar."""
        try:
            if pygame.mixer.music.get_busy():
                pygame.mixer.music.stop()
            pygame.mixer.music.load(file_path)
            pygame.mixer.music.play()
        except Exception as e:
            messagebox.showerror("Hata", f"Ses oynatılamadı: {e}")

    def _speak_last_response(self):
        """Son AI yanıtını OpenAI TTS kullanarak seslendirir (Hyper-realistic)."""
        if not self.last_analysis:
            messagebox.showwarning("Uyarı", "Seslendirilecek bir yanıt yok.")
            return
            
        threading.Thread(target=self._tts_worker, daemon=True).start()

    def _tts_worker(self):
        """TTS işlemini arka planda yapar."""
        try:
            # --- ELEVENLABS SES KLONLAMA KONTROLÜ ---
            if self.eleven_enable_var.get() and self.eleven_manager:
                selected_voice_name = self.eleven_voice_combo.get()
                voice_id = next((v[1] for v in self.eleven_voices if v[0] == selected_voice_name), None)
                
                if voice_id:
                    try:
                        self.animator.start_loading("ElevenLabs Ses Sentezleniyor")
                        temp_mp3 = self.eleven_manager.generate_speech(self.last_analysis[:1000], voice_id)
                        self.animator.stop("Ses Sentezlendi")
                        if temp_mp3:
                            self._play_audio(temp_mp3)
                            return
                    except Exception as e:
                        print(f"ElevenLabs Hatası (Kotanız dolmuş olabilir, OpenAI'a geçiliyor): {e}")
                        # Devam et ve OpenAI TTS'i kullan

            # --- STANDART OPENAI TTS ---
            if not self.api_key:
                self.after(0, lambda: messagebox.showerror("Hata", "OpenAI API anahtarı bulunamadı."))
                return

            from openai import OpenAI
            client = OpenAI(api_key=self.api_key)
            
            # Kullanıcının seçtiği sesi al
            selected_voice_name = self.tts_voice_combo.get()
            selected_voice = self.tts_voices.get(selected_voice_name, "onyx")

            # Text-to-Speech İsteği
            response = client.audio.speech.create(
                model="tts-1",
                voice=selected_voice,
                input=self.last_analysis[:4000]
            )
            
            import time
            temp_tts = f"temp_tts_{int(time.time())}.mp3"
            response.stream_to_file(temp_tts)
            self._play_audio(temp_tts)
        except Exception as e:
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("TTS Hatası", f"Seslendirme başarısız: {err}"))
            if hasattr(self, 'animator'): self.animator.stop("TTS Hatası")

    def _send_quick_chat(self, prompt):
        """Hızlı aksiyon butonları için prompt gönderir."""
        self.chat_entry.delete(0, "end")
        self.chat_entry.insert(0, prompt)
        self.ask_ai_question()

    def load_history_file(self, filename):
        path = os.path.join("recordings", filename)
        self.select_frame_by_name("home")
        threading.Thread(target=lambda: self._transcribe_file(path), daemon=True).start()


    def change_mic(self, value):
        """Kullanıcının seçtiği mikrofon indeksini günceller."""
        try:
            self.selected_mic_index = int(value.split(":")[0])
        except:
            pass

    def _toggle_auto_vad(self):
        """Auto-VAD özelliğini açıp kapatır."""
        self.auto_vad_enabled = self.auto_vad_var.get()
        status = "açıldı" if self.auto_vad_enabled else "kapatıldı"
        print(f"Auto-VAD {status}.")
    def toggle_recording(self, source="home"):
        """
        Kayıt butonuna her basışımızda bu fonksiyon tetiklenir.
        source: "home" veya "language" - Kaydın nereden başlatıldığını belirtir.
        """
        if not self.is_recording:
            self.is_recording = True
            self.active_recording_source = source
            
            # Sadece ilgili butonu güncelle
            btn = self.record_btn if source == "home" else self.coach_record_btn
            btn.configure(text="KAYDI DURDUR", fg_color="red")
            
            self.animator.start_pulse() # Animasyonu başlat
            self.status_label.configure(text="Kaydediliyor...")
            self.audio_frames = []
            
            # VAD Durumlarını Sıfırla
            self.silence_start_time = None
            self.recording_start_time = time.time() # Kayıt başlangıç zamanı
            
            # Çakışmayı önlemek için kayıt işlemini ayrı bir thread'de başlat
            threading.Thread(target=self._record_thread, daemon=True).start()
        else:
            self.is_recording = False
            
            # Sadece aktif olan butonu geri döndür
            btn = self.record_btn if self.active_recording_source == "home" else self.coach_record_btn
            btn.configure(text="KAYDI BAŞLAT", fg_color="green")
            
            self.animator.stop("Kayıt durduruldu.")
            if hasattr(self, 'visualizer'):
                self.visualizer.clear()
            # Asenkron güncellemeyi durduracak bir bayrak gerekirse burada set edilebilir
            # Ancak is_recording False olması yeterli

    def _record_thread(self):
        """Mikrofondan ham ses verilerini okuyan iş parçacığı (Yüksek Öncelikli)."""
        try:
            # Kuyruğu temizle
            while not self.audio_queue.empty():
                self.audio_queue.get()
                
            # Asenkron görselleştirme döngüsünü başlat
            self.after(50, self._update_viz_loop)
            
            # latency='low' ve blocksize=0 (otomatik) ile en kararlı akışı sağla
            with sd.InputStream(samplerate=self.fs, channels=1, callback=self._audio_callback, 
                                device=self.selected_mic_index, blocksize=0, latency='low'):
                while self.is_recording:
                    # Kuyruktan gelen verileri topla
                    try:
                        while not self.audio_queue.empty():
                            data = self.audio_queue.get_nowait()
                            self.audio_frames.append(data)
                            
                            # --- Manuel Kontrol: Kayıt durdurulana kadar devam eder ---
                            # Sadece görselleştirme ve RMS hesaplama (isteğe bağlı) yapılır
                            rms = np.sqrt(np.mean(data**2))
                            self.last_rms = rms
                            
                            # --- Auto-VAD İşlemi (Eğer kullanıcı Ayarlardan açmışsa) ---
                            if self.auto_vad_enabled and (time.time() - self.recording_start_time > 2.0): # 2 sn'den sonra başlasın
                                if rms < self.silence_threshold:
                                    if self.silence_start_time is None:
                                        self.silence_start_time = time.time()
                                    else:
                                        silent_duration = time.time() - self.silence_start_time
                                        if silent_duration > 2.0: # 2 saniye sessizlik yeterli
                                            print(f"Auto-VAD: Sessizlik algılandı ({silent_duration:.1f}s), kayıt durduruluyor.")
                                            self.after(0, self.toggle_recording)
                                            break
                                else:
                                    self.silence_start_time = None
                    except queue.Empty:
                        pass
                    time.sleep(0.05) # İşlemciyi yormadan kuyruğu boşalt
        except Exception as e:
            self.is_recording = False
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("Donanım Hatası", f"Mikrofon hatası: {err}"))
            return
        
        # --- SES İŞLEME: NORMALİZASYON VE GÜRÜLTÜ AZALTMA ---
        if not self.audio_frames:
            self.after(0, lambda: messagebox.showwarning("Kayıt Boş", "Hiç ses verisi alınamadı. Lütfen mikrofonunuzu kontrol edin."))
            return

        try:
            audio_path = "temp_recording.wav"
            audio_data = np.concatenate(self.audio_frames, axis=0)
            
            # 1. Normalizasyon (Ses seviyesini dengeleme)
            max_val = np.max(np.abs(audio_data))
            if max_val > 0:
                audio_data = audio_data / max_val
                
            # 2. Gürültü Azaltma (Eğer aktifse)
                try:
                    # Arka plan gürültüsünü akıllıca azalt (Daha hassas bir oran: 0.6)
                    audio_data = nr.reduce_noise(y=audio_data.flatten(), sr=self.fs, prop_decrease=0.6)
                    audio_data = audio_data.reshape(-1, 1) # Formatı koru
                except Exception as nre:
                    print(f"Gürültü azaltma hatası: {nre}")

            sf.write(audio_path, audio_data, self.fs)
            print(f"Ses işlendi ve kaydedildi: {audio_path}")
        except Exception as e:
            error_msg = str(e)
            self.after(0, lambda msg=error_msg: messagebox.showerror("Ses İşleme Hatası", f"Ses verisi işlenirken hata oluştu: {msg}"))
            return

        # Eğer otomatik kayıt açıksa recordings klasörüne tarih-saat ile kaydet
        if self.autosave_var.get():
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            save_path = os.path.join("recordings", f"kayit_{timestamp}.wav")
            sf.write(save_path, audio_data, self.fs)
            print(f"Ses kaydedildi: {save_path}")

        # Transkripsiyon sürecini başlat
        self._transcribe_file(audio_path)

    def _audio_callback(self, indata, frames, time, status):
        """Mikrofondan gelen ses paketini en hızlı şekilde kuyruğa atar."""
        if status:
            print(f"Ses Akış Durumu: {status}")
        if self.is_recording:
            # Sadece veriyi kopyalayıp kuyruğa at, UI veya Liste işlemi YAPMA!
            self.audio_queue.put(indata.copy())
            self.last_audio_block = indata.copy() # Görselleştirici için son bloğu sakla

    def _update_viz_loop(self):
        """Görselleştiriciyi ana thread üzerinden (asenkron) güncelleyen döngü."""
        if self.is_recording:
            if hasattr(self, 'visualizer') and hasattr(self, 'last_audio_block'):
                self.visualizer.update_visuals(self.last_audio_block)
            # 30ms sonra tekrar çalış (yaklaşık 33 FPS)
            self.after(30, self._update_viz_loop)

    def _transcribe_file(self, path):
        """Ses dosyasını Whisper kullanarak metne dönüştürür."""
        try:
            task = "translate" if self.translate_var.get() else "transcribe"
            model_type = self.model_combo.get()
            
            # Model yükleme veya önbellekten alma
            if self.whisper_model is None or self.current_model_type != model_type:
                self.animator.start_loading(f"Model yükleniyor ({model_type})")
                self.whisper_model = whisper.load_model(model_type, device=self.device)
                self.current_model_type = model_type
            
            self.animator.start_loading("Metne dönüştürülüyor")
            
            # Dil eşleştirmesini yap
            selected_lang_tr = self.lang_combo.get()
            whisper_lang = self.lang_options.get(selected_lang_tr) # None olabilir (auto)
            
            # Whisper transkripsiyon işlemi (En yüksek kalite parametreleri ile)
            res = self.whisper_model.transcribe(
                path, 
                language=whisper_lang, 
                task=task,
                beam_size=5,
                temperature=0.0,
                fp16=True if self.device == "cuda" else False
            )
            
            full_text = res['text'].encode('utf-8', 'replace').decode('utf-8')
            self.last_transcript = full_text 
            self.all_session_transcripts.append({
                "time": datetime.datetime.now().strftime("%H:%M:%S"),
                "text": full_text
            })
            
            # Kaynağa göre ilgili metin kutusuna yazdır
            if self.active_recording_source == "home":
                self.after(0, lambda: self.textbox.insert("end", f"\n[TRANSKRIPT]:\n{full_text}\n"))
                self.after(0, lambda: self.textbox.see("end"))
            elif self.active_recording_source == "language":
                self.after(0, lambda: self.language_textbox.insert("end", f"\n[TRANSKRIPT]:\n{full_text}\n"))
                self.after(0, lambda: self.language_textbox.see("end"))
            elif self.active_recording_source == "topic_chat":
                # Sesle yazma: Metni girişe koy ve otomatik gönder
                self.after(0, lambda: self.topic_chat_entry.delete(0, "end"))
                self.after(0, lambda: self.topic_chat_entry.insert(0, full_text))
                self.after(0, lambda: self.run_topic_ai_chat())
            elif self.active_recording_source == "pronunciation":
                self.after(0, lambda: self._compare_pronunciation(full_text))
                
            # Analiz sekmesi her zaman güncellenebilir (opsiyonel, bağımsızlık için kaldırılabilir)
            self.after(0, lambda: self.analysis_textbox.insert("end", f"\n[TRANSKRIPT]:\n{full_text}\n"))
            
            # İstatistikleri güncelle
            words = len(full_text.split())
            self.stats_manager.add_session(words=words, minutes=0.5) # Yaklaşık 0.5 dk varsayılan çalışma
            self.update_stats_ui()
            
            self.animator.stop("İşlem tamamlandı.")
        except Exception as e:
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("Hata", f"Transkripsiyon Hatası: {err}"))

    def process_audio_file(self):
        """Bilgisayardan bir ses dosyası seçilmesini sağlar."""
        path = filedialog.askopenfilename(filetypes=[("Ses Dosyası", "*.wav *.mp3 *.m4a")])
        if path:
            threading.Thread(target=lambda: self._transcribe_file(path), daemon=True).start()

    def open_recordings_folder(self):
        """Kayıtların tutulduğu klasörü Windows Explorer'da açar."""
        folder_path = os.path.abspath("recordings")
        if os.path.exists(folder_path):
            os.startfile(folder_path)
        else:
            messagebox.showerror("Hata", "Kayıtlar klasörü bulunamadı!")

    # --- GPT-4o ANALİZ METOTLARI ---
    def run_analysis(self):
        """Metin kutusundaki verileri GPT-4o ile analiz etmek üzere gönderir."""
        # Session geçmişini kullanarak zaman damgalı metin oluştur
        text_with_timestamps = ""
        for entry in self.all_session_transcripts:
            text_with_timestamps += f"[{entry['time']}] {entry['text']}\n"
        
        # Eğer geçmiş boşsa (manuel düzeltme yapılmış olabilir), kutudaki ham metni al
        if not text_with_timestamps:
            text_with_timestamps = self.textbox.get("1.0", "end").strip()
            
        if text_with_timestamps:
            threading.Thread(target=self._gpt_logic, args=(text_with_timestamps,), daemon=True).start()

    def _gpt_logic(self, text):
        """Arka planda OpenAI API isteğini yönetir."""
        try:
            if not self.api_key:
                self.after(0, lambda: messagebox.showwarning("Anahtar Eksik", "Lütfen OpenAI API anahtarınızı kaydedin."))
                return

            safe_text = text.encode('utf-8', 'replace').decode('utf-8')
            client = OpenAI(api_key=self.api_key)
            self.animator.start_loading("GPT-4o analiz ediyor")
            
            prompt = self._get_analysis_prompt(safe_text)
            system_msg = self._get_system_prompt()

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ]
            )
            analysis = response.choices[0].message.content
            self._process_analysis_result(analysis, safe_text, "OpenAI")
        except Exception as e:
            err = str(e).encode('utf-8', 'ignore').decode('utf-8')
            self.after(0, lambda err=err: messagebox.showerror("API Hatası", f"Hata: {err}"))

    # --- GEMINI ANALİZ METOTLARI ---
    def run_gemini_analysis(self):
        """Metin kutusundaki verileri Google Gemini ile analiz eder."""
        # Session geçmişini kullanarak zaman damgalı metin oluştur
        text_with_timestamps = ""
        for entry in self.all_session_transcripts:
            text_with_timestamps += f"[{entry['time']}] {entry['text']}\n"
            
        if not text_with_timestamps:
            text_with_timestamps = self.textbox.get("1.0", "end").strip()
            
        if text_with_timestamps:
            threading.Thread(target=self._gemini_logic, args=(text_with_timestamps,), daemon=True).start()

    def _gemini_logic(self, text):
        """Arka planda Gemini API isteğini yönetir."""
        try:
            if not self.gemini_api_key:
                self.after(0, lambda: messagebox.showwarning("Anahtar Eksik", "Lütfen Gemini API anahtarınızı kaydedin."))
                return

            safe_text = text.encode('utf-8', 'replace').decode('utf-8')
            client = GeminiClient(api_key=self.gemini_api_key)
            self.animator.start_loading("Gemini analiz ediyor")
            
            prompt = self._get_analysis_prompt(safe_text)
            system_msg = self._get_system_prompt()
            
            analysis = client.generate_content(prompt, system_instruction=system_msg)
            self._process_analysis_result(analysis, safe_text, "Gemini")
        except Exception as e:
            err = str(e).encode('utf-8', 'ignore').decode('utf-8')
            self.after(0, lambda err=err: messagebox.showerror("API Hatası", f"Hata: {err}"))

    # --- AI CHAT (SORU-CEVAP) MANTIĞI ---
    def ask_ai_question(self):
        """Kullanıcının sorusunu transkript ile birlikte AI'ya gönderir."""
        question = self.chat_entry.get().strip()
        transcript = self.textbox.get("1.0", "end").strip()
        
        if not question: return
        if not transcript:
            messagebox.showwarning("Uyarı", "Önce bir ses kaydı veya dosya yüklemelisin.")
            return
            
        self.ask_btn.configure(state="disabled", text="...")
        threading.Thread(target=self._chat_logic, args=(question, transcript), daemon=True).start()

    def _chat_logic(self, question, transcript):
        """Arka planda AI chat isteğini yönetir."""
        try:
            # Varsa Gemini, yoksa OpenAI kullan
            system_msg = self._get_system_prompt()
            if self.gemini_api_key:
                client = GeminiClient(api_key=self.gemini_api_key)
                prompt = f"Şu transkript üzerinden soruyu cevapla:\n\nTRANSKRİPT:\n{transcript}\n\nSORU: {question}"
                response = client.generate_content(prompt, system_instruction=system_msg)
                answer = response
            elif self.api_key:
                client = OpenAI(api_key=self.api_key)
                prompt = f"Şu transkript üzerinden soruyu cevapla:\n\nTRANSKRİPT:\n{transcript}\n\nSORU: {question}"
                res = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_msg},
                        {"role": "user", "content": prompt}
                    ]
                )
                answer = res.choices[0].message.content
            else:
                self.after(0, lambda: messagebox.showwarning("Hata", "Lütfen API anahtarlarını kontrol et."))
                return

            self.last_analysis = answer # Seslendirilebilmesi için son cevabı kaydet
            self.after(0, lambda q=question, a=answer: self._add_chat_to_ui(q, a))
        except Exception as e:
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("Chat Hatası", f"Hata: {err}"))
        finally:
            self.after(0, lambda: self.ask_btn.configure(state="normal", text="SOR"))
            self.after(0, lambda: self.chat_entry.delete(0, "end"))

    def _add_chat_to_ui(self, question, answer):
        """Soruyu ve cevabı analiz kutusuna ekler."""
        chat_text = f"\n\n--- SORU-CEVAP ---\nSoru: {question}\nCevap: {answer}\n------------------\n"
        self.analysis_textbox.insert("end", chat_text)
        self.analysis_textbox.see("end")
        self.status_label.configure(text="AI sorunu cevapladı.")

    def _get_analysis_prompt(self, safe_text):
        """AI modellerine akademik ve profesyonel bitirme projesi seviyesinde analiz komutu döner."""
        return f"""
        GÖREV: Aşağıdaki zaman damgalı transkriptleri analiz et.
        
        [KRİTİK TALİMATLAR]:
        1. HER BİR transkript segmentini (zaman damgasıyla birlikte) MUTLAKA ayrı ayrı incele.
        2. Raporun her alt başlığında hangi segmentten bahsettiğini KÖŞELİ PARANTEZ içindeki zaman damgasıyla BELİRT (Örn: '[12:45:00] kaydında...', '[13:00:10] segmenti gösteriyor ki...').
        3. Zaman damgalarını asla atlama, her paragrafın başında veya sonunda hangi kayda ait olduğu yazılsın.
        4. Analiz sonucunda skorları ve segmentlerin duygu durumlarını aşağıda istenen formatta sağla.
        
        RAPOR FORMATI:
        
        1. YÖNETİCİ ÖZETİ (Executive Summary):
           - Konuşmanın ana amacını, bağlamını ve en önemli sonucunu 4-5 cümlelik akademik bir dille özetle.
        
        2. DETAYLI KONU VE İÇERİK ANALİZİ:
           - Kayıtta geçen temel temaları, kavramları ve tartışılan konuları derinlemesine analiz et. 
           - Varsa teknik terimleri ve bunların bağlam içindeki kullanımını açıkla.
        
        3. STRATEJİK BULGULAR VE ANALİZ:
           - Konuşmanın arka planındaki stratejik hedefleri veya temel mesajları saptayın.
           - Konuşmacıların argümanlarını ve fikir birliği/ayrılığı noktalarını belirtin.
        
        4. DERİN DUYGU VE TONLAMA ANALİZİ:
           - Metnin genel duygusal haritasını çıkar (Örn: Heyecanlı, Kaygılı, Profesyonel, Çözüm Odaklı).
           - Bu tonlamanın konuşmanın amacına etkisini yorumla.
        
        5. AKSİYON MADDELERİ VE EYLEM PLANI:
           - Konuşmada belirlenen görevleri, sorumlulukları ve atılması gereken adımları liste formatında (Bullet Points) yaz.
        
        6. AKADEMİK SONUÇ VE ÖNERİLER:
           - Analiz edilen verilere dayanarak, gelecekte yapılabilecek geliştirmeler veya iyileştirmeler için profesyonel tavsiyeler sun.
        
        [[DATA_START]]
        (ÖNEMLİ: Bu satırdan sonrasını SADECE veri formatında hazırla. Kullanıcı bu kısmı görmeyecek.)
        
        POZİTİF: [sayı]
        NEGATİF: [sayı]
        NÖTR: [sayı]
        
        SEGMENTS:
        [
          {{"text": "...", "sentiment": "pos/neg/neu"}},
          ...
        ]
        
        (ÖNEMLİ: Zaman çizelgesi için metni küçük parçalara/cümlelere böl ve duygusunu KESİN JSON formatında sağla. JSON bloğunda anahtar ve değerler için çift tırnak (") kullan.)
        SEGMENTS:
        [
          {{"text": "...", "sentiment": "pos/neg/neu"}},
          ...
        ]
        
        [ANALİZ EDİLECEK METİN]:
        {safe_text}
        """

    # --- AI DİL KOÇU MANTIĞI ---
    def run_language_analysis(self):
        """
        Dil Koçu sekmesinde girdiğimiz metni veya transkripti parçalayarak
        öğrenciye özel akademik geri bildirim oluşturmak için bu fonksiyonu çağırıyoruz.
        """
        text = self.textbox.get("1.0", "end").strip()
        if not text:
            # Eğer dashboard boşsa kendi kutusuna bak
            text = self.language_textbox.get("1.0", "end").replace("--- AI DİL KOÇU HAZIR ---\nLütfen bir ses kaydı yapın veya metin girin, ardından 'DİL ANALİZİ BAŞLAT' butonuna basın.\n", "").strip()
        
        if not text:
            messagebox.showwarning("Uyarı", "Analiz edilecek metin yok!")
            return
            
        self.language_analysis_result = ""
        self.coach_chat_history = [] # Yeni analizde geçmişi sıfırla
        self.run_coach_btn.configure(state="disabled", text="ANALİZ EDİLİYOR...")
        threading.Thread(target=self._language_coach_logic, args=(text,), daemon=True).start()

    def _language_coach_logic(self, text):
        """Arka planda Dil Koçu API isteğini yönetir."""
        try:
            # Varsa Gemini, yoksa OpenAI kullan
            target_lang = self.coach_lang_combo.get()
            level = self.coach_level_combo.get()
            mode = self.coach_mode_combo.get()
            
            self.animator.start_loading(f"Dil Koçu ({target_lang}) analiz ediyor")
            
            prompt = self._get_language_coach_prompt(text, target_lang, level, mode)
            system_msg = "Sen uzman bir dil eğitmeni ve polyglot bir mentorsun. Öğrencilerine destekleyici, öğretici ve profesyonel geri bildirimler verirsin."

            if self.gemini_api_key:
                client = GeminiClient(api_key=self.gemini_api_key)
                response = client.generate_content(prompt, system_instruction=system_msg)
                result = response
            elif self.api_key:
                client = OpenAI(api_key=self.api_key)
                res = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_msg},
                        {"role": "user", "content": prompt}
                    ]
                )
                result = res.choices[0].message.content
            else:
                self.after(0, lambda: messagebox.showwarning("Hata", "Lütfen API anahtarlarını kontrol et."))
                return

            self.language_analysis_result = result
            self.after(0, lambda r=result: self._update_language_ui(r))
        except Exception as e:
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("Dil Koçu Hatası", f"Hata: {err}"))
        finally:
            self.after(0, lambda: self.run_coach_btn.configure(state="normal", text="DİL ANALİZİ BAŞLAT"))

    # --- SENARYO YÖNETİMİ ---
    def _on_topic_change(self, choice):
        """Konu değiştiğinde senaryo listesini güncelle."""
        if choice in self.scenarios_data:
            scenarios = list(self.scenarios_data[choice].keys())
            self.scenario_combo.configure(values=scenarios)
            if scenarios:
                self.scenario_combo.set(scenarios[0])
                self._on_scenario_change(scenarios[0])
            else:
                self.scenario_combo.set("Senaryo Yok")
                self.sub_option_combo.configure(values=[])
                self.sub_option_combo.set("-")

    def _on_scenario_change(self, choice):
        """Senaryo değiştiğinde alt seçenekleri (karakterleri) güncelle."""
        topic = self.topic_combo.get()
        if topic in self.scenarios_data and choice in self.scenarios_data[topic]:
            sub_options = self.scenarios_data[topic][choice]
            if sub_options:
                self.sub_option_combo.configure(state="normal", values=sub_options)
                self.sub_option_combo.set(sub_options[0])
            else:
                self.sub_option_combo.configure(values=[], state="disabled")
                self.sub_option_combo.set("-")
        
        # Ambiyans Sesini Güncelle
        if hasattr(self, 'sound_manager'):
            self.sound_manager.play_ambience(choice)
            
        # Otomatik Görsel Üretimini Tetikle
        self.manual_image_generation()

    def create_custom_scenario(self):
        """Kullanıcıdan bir ilgi alanı alıp AI'ya özel senaryo ürettirir."""
        dialog = ctk.CTkInputDialog(text="Hangi konuda bir senaryo oluşturmak istersin?\n(Örn: Uzay Mimarisi, Robotik Cerrahi, vb.)", title="Akıllı Senaryo Jeneratörü")
        interest = dialog.get_input()
        
        if interest:
            if not self.api_key and not self.gemini_api_key:
                messagebox.showerror("Hata", "Lütfen önce 'Ayarlar' sekmesinden bir API anahtarı kaydedin.")
                return
            
            self.magic_wand_btn.configure(state="disabled", text="✨")
            threading.Thread(target=self._custom_scenario_logic, args=(interest,), daemon=True).start()

    def _custom_scenario_logic(self, interest):
        """AI'dan ilgi alanına uygun senaryo ve karakterler üretir."""
        try:
            print(f"[*] Özel senaryo üretiliyor: {interest}")
            prompt = f"Kullanıcı '{interest}' konusunda bir dil eğitimi/RPG senaryosu istiyor. " \
                     f"Lütfen bir senaryo adı, kısa bir açıklama ve 3 adet farklı karakter/mod adı üret. " \
                     f"Format: SADECE JSON döndür. Örnek: {{\"scenario_name\": \"...\", \"description\": \"...\", \"characters\": [\"...\", \"...\", \"...\"]}} " \
                     f"Dili Türkçe olsun. JSON tırnakları için ÇİFT TIRNAK kullan."
            
            system_msg = "Sen yaratıcı bir eğitim tasarımcısısın. SADECE JSON formatında yanıt ver, başka açıklama ekleme."
            
            result = ""
            used_model = "None"
            # 1. Öncelik: Gemini
            if self.gemini_api_key:
                try:
                    client = GeminiClient(api_key=self.gemini_api_key)
                    response = client.generate_content(prompt, system_instruction=system_msg)
                    # Eğer hata mesajı DEĞİLSE ve boş değilse sonucu al
                    if response and not response.startswith("[V19]"):
                        result = response
                        used_model = "Gemini"
                    else:
                        print(f"Gemini Kota/Hata (Fallback): {response[:100]}...")
                except Exception as gem_ex:
                    print(f"Gemini Bağlantı Hatası: {gem_ex}")

            # 2. Öncelik/Fallback: OpenAI
            if not result and self.api_key:
                try:
                    print("[*] GPT-4o ile devam ediliyor...")
                    client = OpenAI(api_key=self.api_key)
                    res = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "system", "content": system_msg}, {"role": "user", "content": prompt}]
                    )
                    result = res.choices[0].message.content
                    used_model = "GPT-4o"
                except Exception as gpt_ex:
                    print(f"GPT Hatası: {gpt_ex}")
            
            if result:
                import json
                import re
                
                # Daha sağlam JSON ayıklama (Regex kullanarak)
                json_match = re.search(r'\{.*\}', result, re.DOTALL)
                if json_match:
                    clean_json = json_match.group(0)
                else:
                    clean_json = result.replace("```json", "").replace("```", "").strip()
                
                data = json.loads(clean_json)
                
                s_name = data.get("scenario_name", interest)
                s_desc = data.get("description", "")
                chars = data.get("characters", ["Uzman", "Öğrenci", "Mentor"])
                
                # Mevcut verilere ekle
                custom_topic = "Özel Senaryo"
                if custom_topic not in self.scenarios_data:
                    self.scenarios_data[custom_topic] = {}
                
                self.scenarios_data[custom_topic][s_name] = chars
                
                # UI Güncelle
                self.after(0, lambda t=custom_topic, s=s_name: self._update_custom_topic_ui(t, s))
            
            else:
                # Eğer hiçbir modelden sonuç alınamadıysa
                self.after(0, lambda: messagebox.showwarning("Hata", "AI senaryo üretemedi. Lütfen API kotalarını veya bağlantınızı kontrol edin."))
                
        except Exception as e:
            error_msg = str(e)
            self.after(0, lambda m=error_msg: messagebox.showerror("Hata", f"Senaryo oluşturulamadı: {m}"))
        finally:
            self.after(0, lambda: self.magic_wand_btn.configure(state="normal", text="🪄"))

    def _update_custom_topic_ui(self, topic, scenario):
        """Yeni üretilen senaryoyu combo boxlara ekler."""
        # Konu listesini güncelle
        current_topics = list(self.scenarios_data.keys())
        self.topic_combo.configure(values=current_topics)
        self.topic_combo.set(topic)
        
        # Senaryo listesini güncelle
        scenarios = list(self.scenarios_data[topic].keys())
        self.scenario_combo.configure(values=scenarios)
        self.scenario_combo.set(scenario)
        
        # Karakterleri yükle
        sub_options = self.scenarios_data[topic][scenario]
        self.sub_option_combo.configure(values=sub_options, state="normal")
        self.sub_option_combo.set(sub_options[0])
        
        messagebox.showinfo("Başarılı", f"'{scenario}' senaryosu oluşturuldu!\nSohbet başlatılıyor...")
        
        # Seçenekleri tetikle ve sohbeti başlat
        self._on_scenario_change(scenario)
        self.run_topic_ai_chat()

    def show_word_bank(self):
        """Kelime Defteri penceresini açar."""
        wb_window = ctk.CTkToplevel(self)
        wb_window.title("Kelime Defterim")
        wb_window.geometry("500x600")
        wb_window.attributes("-topmost", True)
        
        ctk.CTkLabel(wb_window, text="📔 KAYDEDİLEN KELİMELER", font=("Inter", 18, "bold")).pack(pady=10)
        
        wb_textbox = ctk.CTkTextbox(wb_window, font=("Inter", 14), width=450, height=450)
        wb_textbox.pack(pady=10, padx=10)
        
        # Kelimeleri Yükle
        words = self.load_words()
        if not words:
            wb_textbox.insert("1.0", "Henüz kelime kaydedilmemiş.\n\nİpucu: Dil Koçu sekmesinde analiz sonucunda çıkan kelimeleri buraya ekleyebilirsiniz.")
        else:
            for word, data in words.items():
                wb_textbox.insert("end", f"📌 {word.upper()}:\n   Anlam: {data.get('meaning', '-')}\n   Örnek: {data.get('example', '-')}\n{'-'*40}\n")
        
        wb_textbox.configure(state="disabled")
        
        ctk.CTkButton(wb_window, text="KAPAT", command=wb_window.destroy).pack(pady=10)

    def load_words(self):
        """JSON dosyasından kelimeleri yükler."""
        path = "word_bank.json"
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        return {}

    def save_word(self, word, meaning, example):
        """Yeni bir kelimeyi JSON dosyasına kaydeder."""
        words = self.load_words()
        words[word.lower()] = {"meaning": meaning, "example": example, "date": str(datetime.datetime.now())}
        with open("word_bank.json", "w", encoding="utf-8") as f:
            json.dump(words, f, ensure_ascii=False, indent=4)
        messagebox.showinfo("Başarılı", f"'{word}' kelime defterine eklendi!")

    def start_pronunciation_test(self):
        """AI'nın son önerisini kullanıcının tekrar etmesini ister ve karşılaştırır."""
        if not self.last_transcript:
            messagebox.showwarning("Uyarı", "Önce bir analiz yapmalısınız ki karşılaştıracak bir cümle olsun.")
            return
            
        # Basit bir regex ile önerilen cümleyi bulmaya çalış (Daha gelişmiş olabilir)
        # Genelde 'Öneri:' veya 'Suggestion:' sonrası cümleyi alabiliriz.
        import re
        suggestion = ""
        lines = self.language_analysis_result.split("\n")
        for line in lines:
            if "➤" in line or "Öneri:" in line or "Suggestion:" in line:
                suggestion = line.split(":")[-1].strip().replace("➤", "").strip()
                break
        
        if not suggestion:
            # Eğer özel bir öneri bulunamazsa tüm transkripti veya son parçayı al
            suggestion = self.last_transcript.strip()

        msg = f"Lütfen şu cümleyi yüksek sesle tekrar et:\n\n\"{suggestion}\"\n\nKayıt otomatik başlayacak."
        if messagebox.askokcancel("Telaffuz Testi", msg):
            self.target_test_sentence = suggestion
            self.active_recording_source = "pronunciation"
            self.toggle_recording(source="language") # Kaydı başlat

    def _compare_pronunciation(self, user_text):
        """Kullanıcının söylediği ile hedef cümleyi karşılaştırır."""
        target = self.target_test_sentence.lower().strip()
        user = user_text.lower().strip()
        
        # Gereksiz noktalamaları temizle
        import string
        target = target.translate(str.maketrans('', '', string.punctuation))
        user = user.translate(str.maketrans('', '', string.punctuation))
        
        # Kelime bazlı karşılaştırma
        target_words = target.split()
        user_words = user.split()
        
        matches = 0
        for word in user_words:
            if word in target_words:
                matches += 1
                
        score = int((matches / len(target_words)) * 100) if target_words else 0
        
        result_msg = f"Hedef: {self.target_test_sentence}\nSöylenen: {user_text}\n\n"
        result_msg += f"🎯 Telaffuz Skoru: %{score}\n\n"
        
        if score > 90:
            result_msg += "Mükemmel! Tıpkı bir ana dil konuşuru gibisin. 🌟"
        elif score > 60:
            result_msg += "Gayet iyi, birkaç kelime üzerinde durabilirsin. 👍"
        else:
            result_msg += "Biraz daha pratik yapmalısın. Pes etme! 💪"
            
        messagebox.showinfo("Telaffuz Sonucu", result_msg)
        self.active_recording_source = "language" # Eski haline dön

    # --- KONU BAZLI AI SOHBET MANTIĞI ---
    def run_topic_ai_chat(self):
        """Seçili konu üzerinden AI ile bağımsız bir sohbet başlatır veya devam ettirir."""
        topic = self.topic_combo.get()
        scenario = self.scenario_combo.get()
        sub_option = self.sub_option_combo.get()
        user_input = self.topic_chat_entry.get().strip()
        
        # Eğer start butonuna basıldıysa ama input boşsa, başlatma mesajı iste
        is_start = False
        if not user_input:
            is_start = True
            # Sohbet alanını temizle ve başlangıç mesajını göster
            self.topic_textbox.delete("1.0", "end")
            self.topic_textbox.insert("end", f"--- {sub_option if sub_option != '-' else scenario} ile Bağlantı Kuruluyor... ---\n")
            
            # Yeni bir sohbet başlatılıyorsa görseli güncelle
            threading.Thread(target=self.generate_topic_image, args=(scenario, f"{scenario} - {sub_option} context"), daemon=True).start()
        
        self.start_topic_btn.configure(state="disabled", text="...")
        self.topic_ask_btn.configure(state="disabled")
        
        threading.Thread(target=self._topic_chat_logic, args=(topic, user_input, scenario, sub_option), daemon=True).start()

    def generate_topic_image(self, topic, description):
        """DALL-E 3 kullanarak konuya uygun görsel oluşturur ve arayüze basar."""
        if not self.api_key:
            return

        try:
            # UI işlemleri ana thread'de yapılmalı
            self.after(0, lambda: self.topic_image_label.configure(text="Görsel Oluşturuluyor..."))
            
            client = OpenAI(api_key=self.api_key)
            
            # Prompt'u optimize et ve güvenlik filtreleri için rafine et
            # DALL-E'nin 'safe' politikalarına uygun bir dille betimleme yap
            bad_words = ["vahşet", "kan", "savaş", "ölüm", "şiddet", "silah", "saldırı", "katliam", "intikam", "kılıç", "ok", "kalkan", "yaralı", "ceset"]
            safe_description = description.lower()
            for word in bad_words:
                safe_description = safe_description.replace(word, "tarihi atmosfer")
            
            image_prompt = f"Educational and atmospheric concept art: {topic}. {safe_description}. Detailed, 4k, cinematic, oil painting style."
            
            # Tarihsel/Kültürel Bağlam Güçlendirme (Örn: Osmanlı/İstanbul)
            historical_keywords = ["İstanbul'un Fethi", "Osmanlı", "Byzantine", "Ottoman", "Constantinople"]
            if any(key.lower() in safe_description.lower() or key.lower() in topic.lower() for key in historical_keywords):
                image_prompt = f"Historical accurate 15th century Ottoman/Byzantine atmosphere. {safe_description}. " \
                               f"Minarets, ancient city walls, historical architecture, historical ships. " \
                               f"Cinematic lighting, detailed, 4k, no modern elements."

            if topic == "RPG Oyunu":
                image_prompt = f"Atmospheric, immersive concept art. Scene: {safe_description}. " \
                               f"High detail, cinematic lighting, 4k."
                if any(key.lower() in safe_description.lower() for key in historical_keywords):
                    image_prompt += " Ottoman Empire architecture and 15th century historical style."

            try:
                response = client.images.generate(
                    model="dall-e-3",
                    prompt=image_prompt[:1000],
                    size="1792x1024",
                    quality="standard",
                    n=1,
                )
            except Exception as e:
                err_msg = str(e).lower()
                if "content_policy_violation" in err_msg or "safety_system" in err_msg:
                    # İlk deneme filtreye takıldıysa daha güvenli bir dille sessizce tekrar dene
                    safe_image_prompt = f"Peaceful and educational concept art illustration for {topic}. Cinematic lighting, soft colors, professional concept art."
                    try:
                        response = client.images.generate(
                            model="dall-e-3",
                            prompt=safe_image_prompt,
                            size="1792x1024",
                            quality="standard",
                            n=1,
                        )
                    except Exception as e2:
                        raise Exception(f"Görsel üretimi güvenlik kısıtlamasına takıldı: {e2}")
                else:
                    raise e

            image_url = response.data[0].url
            
            # Resmi indir
            response = requests.get(image_url)
            img_data = response.content
            
            # PIL ile aç ve CTkImage'a çevir
            pil_image = Image.open(io.BytesIO(img_data))
            
            # Frame boyutuna göre oranla (Height 300 sabit - Geniş ölçek)
            aspect_ratio = pil_image.width / pil_image.height
            new_height = 300
            new_width = int(new_height * aspect_ratio)
            
            ctk_image = ctk.CTkImage(light_image=pil_image, dark_image=pil_image, size=(new_width, new_height))
            
            self.after(0, lambda: self.topic_image_label.configure(image=ctk_image, text=""))
            
        except Exception as e:
            err_msg = str(e)
            print(f"Görsel oluşturma hatası: {err_msg}")
            # Güvenlik hatası ise kullanıcıyı bilgilendir
            if "content_policy_violation" in err_msg:
                self.after(0, lambda: self.topic_image_label.configure(text="Güvenlik Politikası Gereği Görsel Üretilemedi"))
            else:
                self.after(0, lambda: self.topic_image_label.configure(text="Sistem Hatası: Görsel Üretilemedi"))

    def manual_image_generation(self):
        """Kullanıcının isteğiyle görsel oluşturur."""
        topic = self.topic_combo.get()
        scenario = self.scenario_combo.get()
        
        desc = f"{topic} - {scenario}"
        if hasattr(self, 'last_topic_response') and self.last_topic_response:
             desc += f". Context: {self.last_topic_response[:150]}"
        
        threading.Thread(target=self.generate_topic_image, args=(topic, desc), daemon=True).start()

    def upload_topic_notes(self):
        """Kullanıcının yüklediği ders notlarını (PDF/TXT) okur ve bağlama ekler."""
        file_path = filedialog.askopenfilename(filetypes=[("Belgeler", "*.pdf *.txt")])
        if not file_path:
            return

        try:
            content = ""
            if file_path.lower().endswith('.pdf'):
                # PyPDF2 veya benzeri bir kütüphane gerekebilir ama şimdilik basit text parse deneyelim
                # Eğer yoksa hata verebilir, bu yüzden basitleştirilmiş bir yaklaşım kullanalım veya kullanıcıdan text isteyelim.
                # Project'te PyPDF2 yoksa, pypdf import etmeyi deneyelim
                try:
                    import pypdf
                    reader = pypdf.PdfReader(file_path)
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                except ImportError:
                    # Kullanıcıya bilgi ver
                    messagebox.showinfo("Bilgi", "PDF okumak için 'pypdf' kütüphanesi gerekli. Text dosyası yüklemeyi deneyin.")
                    return
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

            if content:
                # İçeriği sakla
                self.uploaded_notes_contet = content[:5000] # Çok uzun olmasın, token limiti
                messagebox.showinfo("Başarılı", "Notlar yüklendi! Artık sorularınızı bu notlara göre sorabilirsiniz.")
                self.topic_textbox.insert("end", f"\n[SİSTEM]: '{os.path.basename(file_path)}' içeriği bağlama eklendi.\n")
        except Exception as e:
            messagebox.showerror("Hata", f"Dosya okunamadı: {e}")

    def _topic_chat_logic(self, topic, user_input, scenario, sub_option):
        """Arka planda bağımsız konu chat isteğini yönetir ve hafızayı kullanır."""
        try:
            # Hafızayı Derle (Son 5 mesaj)
            history_context = ""
            for h in self.topic_chat_history[-5:]:
                history_context += f"Öğrenci: {h['input']}\nSen: {h['output']}\n"

            # Sistem Mesajını (Persona) Oluştur
            system_msg = f"Sen {topic} konusunda uzmansın. "
            
            # Senaryo ve Karakter Entegrasyonu
            if scenario and scenario != "Senaryo Seçiniz":
                system_msg += f"Şu anki modun: '{scenario}'. "
                # Karakter varsa
                if sub_option and sub_option != "-" and sub_option != "Karakter Seçiniz":
                    system_msg += f"Karakterin: '{sub_option}'. "
                    if sub_option in self.character_styles:
                        system_msg += f"TARZ TALİMATI: {self.character_styles[sub_option]} "
                    else:
                        system_msg += f"Lütfen BU KARAKTER GİBİ konuş, onun sözlerini veya tarzını taklit et. "
                
                # Özel Senaryo Talimatları
                if "Sokratik" in scenario:
                    system_msg += "ASLA direkt cevap verme. Sadece soru sorarak öğrencinin bulmasını sağla. "
                elif "Code Review" in scenario:
                    system_msg += "Kodu çok sıkı eleştir, hataları bul, best practice öner. "
                elif "5 Yaşındayım" in scenario or "Feynman" in scenario:
                    system_msg += "Çok basit, analojilerle ve eğlenceli anlat. "
                elif "Münazara" in scenario or "Şeytanın Avukatı" in scenario:
                    system_msg += "Öğrencinin fikrine nazikçe ama zekice karşı çık, antitez sun. "
                elif "Mülakat" in scenario:
                    system_msg += "Mülakat yapıyorsun. Zor teknik sorular sor, cevabı puanla. "
                
                if topic == "RPG Oyunu":
                    system_msg += "Sen bir GM'sın (Oyun Yöneticisi). Hikaye anlat, betimle ve mutlaka 2-3 adet numaralı seçenek sun. "
                    system_msg += "ÖNEMLİ: Seçenekleri MUTLAKA '1. Seçenek Adı' formatında yeni satırlarda yaz. "
                    system_msg += "Cevabının sonunda MUTLAKA oyuncunun canını ve envanterini etiket içinde belirt. "
                    system_msg += "Format: [HP:80] [INV:Kılıç,Meşale] gibi. "
                    system_msg += "Başlangıçta HP:100 ve INV:Boş olsun. Olaylara göre güncelle. "
                
                system_msg += "Cana yakın ve öğretici bir dille yardımcı ol."

            # RAG (Doküman) Entegrasyonu
            # RAG (Doküman) Entegrasyonu - EN YÜKSEK ÖNCELİK
            # Eğer doküman yüklendiyse, sistem mesajının başına ekliyoruz ki talimatları override edebilsin.
            if hasattr(self, 'uploaded_notes_contet') and self.uploaded_notes_contet:
                rag_instruction = f"""
                [KULLANICI EKLİ DOKÜMAN BAŞLANGIÇ]
                {self.uploaded_notes_contet}
                [KULLANICI EKLİ DOKÜMAN BİTİŞ]
                
                TALİMAT: Cevaplarını SADECE ve ÖNCELİKLE yukarıdaki dokümana dayandır. 
                Seçili konu ({topic}) veya senaryo ({scenario}) ne olursa olsun, dokümandaki bilgiyi esas al.
                Eğer sorunun cevabı dokümanda yoksa: "Bu bilgi yüklediğiniz notlarda yer almıyor." de ve genel bilgini kullan.
                """
                # RAG talimatını başa ekle
                system_msg = rag_instruction + "\n" + system_msg

            prompt = self._get_topic_prompt(topic, user_input, scenario, sub_option, history_context)
            
            # RPG Görsel Tetikleyici
            if topic == "RPG Oyunu" and (len(self.topic_chat_history) == 0):
                 threading.Thread(target=self.generate_topic_image, args=(scenario, f"{scenario} atmosphere"), daemon=True).start()

            # API Çağrısı ve Fallback (Yedekleme) Mantığı
            result = ""
            used_model = "None"
            
            # 1. Öncelik: Gemini
            if self.gemini_api_key:
                try:
                    client = GeminiClient(api_key=self.gemini_api_key)
                    response = client.generate_content(prompt, system_instruction=system_msg)
                    
                    # Hata kontrolü: Eğer response bir string ise ve hata mesajı içeriyorsa
                    if isinstance(response, str) and ("quota_dimensions" in response or "RESOURCE_EXHAUSTED" in response or "quota" in response):
                        raise Exception(f"Gemini Quota Error: {response}")
                        
                    if response:
                        result = response
                        used_model = "Gemini"
                except Exception as gemini_error:
                    err = str(gemini_error)
                    print(f"Gemini Hatası (Fallback devreye giriyor): {err}")
                    # Gemini hata verdiyse ve OpenAI key varsa devam et, yoksa hata fırlat
                    if not self.api_key:
                        self.after(0, lambda err=err: messagebox.showerror("API Hatası", f"Gemini hatası ve OpenAI anahtarı yok: {err}"))
                        return

            # 2. Öncelik (veya Fallback): OpenAI (GPT)
            if not result and self.api_key:
                try:
                    client = OpenAI(api_key=self.api_key)
                    res = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {"role": "system", "content": system_msg},
                            {"role": "user", "content": prompt}
                        ]
                    )
                    result = res.choices[0].message.content
                    used_model = "GPT-4o"
                except Exception as gpt_error:
                     err = str(gpt_error)
                     self.after(0, lambda err=err: messagebox.showerror("API Hatası", f"Tüm modeller başarısız oldu. {err}"))
                     return
            
            if not result:
                self.after(0, lambda: messagebox.showwarning("Hata", "API anahtarı eksik veya geçersiz."))
                return

            self.last_topic_response = result
            self.topic_chat_history.append({"topic": topic, "input": user_input, "output": result})
            self.after(0, lambda t=topic, u=user_input, r=result: self._update_topic_ui(t, u, r))
            
            # Otomatik Seslendirme Kontrolü
            if self.auto_tts_topic_var.get():
                self.after(0, self._speak_topic_last_response)
                
        except Exception as e:
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("Konu Sohbet Hatası", f"Hata: {err}"))
        finally:
            self.after(0, lambda: self.start_topic_btn.configure(state="normal", text="SOHBETİ BAŞLAT"))
            self.after(0, lambda: self.topic_ask_btn.configure(state="normal"))
            self.after(0, lambda: self.topic_chat_entry.delete(0, "end"))

    def _get_topic_prompt(self, topic, user_input, scenario, sub_option, history=""):
        prompt = f"Konu: {topic}\nSenaryo: {scenario}\n"
        if sub_option and sub_option != "-":
            prompt += f"Karakterin: {sub_option}\n"
            
        if history:
            prompt += f"--- Sohbet Geçmişi ---\n{history}\n"
        
        if not user_input:
            if sub_option and sub_option != "-":
                prompt += f"GÖREV: Lütfen kendini '{sub_option}' olarak tanıtarak sohbete başla. " \
                          f"Karakterinin ismini ilk cümlede mutlaka kullan (Örn: 'Merhaba, ben {sub_option}...'). " \
                          f"Kendi tarzınla kullanıcıyı selamla ve '{scenario}' senaryosu hakkında kısa bir giriş yapıp ilk sorunla sohbeti başlat."
            else:
                prompt += f"{topic} ve {scenario} hakkında bana ilham verici bir başlangıç bilgisi veya fikir ver."
        else:
            prompt += f"Öğrenci: {user_input}"
        return prompt

    def _speak_topic_last_response(self):
        """Mevcut seanstaki son AI cevabını seslendirir."""
        if not self.last_topic_response:
            return
            
        def tts_worker():
            try:
                self.status_label.configure(text="Ses hazırlanıyor...")
                # --- ELEVENLABS SES KLONLAMA KONTROLÜ ---
                if self.eleven_enable_var.get() and self.eleven_manager:
                    selected_voice_name = self.eleven_voice_combo.get()
                    voice_id = next((v[1] for v in self.eleven_voices if v[0] == selected_voice_name), None)
                    if voice_id:
                        temp_mp3 = self.eleven_manager.generate_speech(self.last_topic_response[:1000], voice_id)
                        if temp_mp3:
                            self.status_label.configure(text="Ses oynatılıyor (ElevenLabs)...")
                            self._play_audio(temp_mp3)
                            return
                
                # --- STANDART OPENAI TTS ---
                self.status_label.configure(text="Ses hazırlanıyor (OpenAI)...")
                # Karakter sesini belirle (Eğer seçilen bir karakter varsa)
                sub_option = self.sub_option_combo.get()
                character_voice = self.character_voices.get(sub_option)
                
                # Eğer karakter sesi varsa onu kullan, yoksa ayarlardaki sesi kullan
                selected_voice = character_voice if character_voice else self.tts_voices.get(self.tts_voice_combo.get(), "nova")
                
                client = OpenAI(api_key=self.api_key)
                response = client.audio.speech.create(
                    model="tts-1",
                    voice=selected_voice,
                    input=self.last_topic_response[:2000] # Hız için limit
                )
                temp_file = f"temp_topic_tts_{int(time.time())}.mp3"
                response.stream_to_file(temp_file)
                self.status_label.configure(text="Ses oynatılıyor...")
                self._play_audio(temp_file)
            except Exception as e:
                self.status_label.configure(text=f"Ses Hatası: {e}")
                print(f"Topic TTS Error: {e}")
                
        threading.Thread(target=tts_worker, daemon=True).start()

    def _update_topic_ui(self, topic, user_input, result):
        if user_input:
            msg = f"\n[SEN]: {user_input}\n[AI ({topic})]: {result}\n"
        else:
            msg = f"\n--- {topic} HAKKINDA BİR FİKİR/ÖNERİ ---\n{result}\n"
            
        self.topic_textbox.insert("end", msg)
        self.topic_textbox.see("end")
        self.status_label.configure(text=f"{topic} sohbeti güncellendi.")
        
        # Eğer RPG Modundaysak ve AI seçenekler sunduysa butonları göster
        if topic == "RPG Oyunu":
            self._update_rpg_stats(result)
            self._parse_and_show_rpg_choices(result)
            # Otomatik Görsel Güncelleme (Her mesajda)
            threading.Thread(target=self.generate_topic_image, args=("RPG Oyunu", result[:200]), daemon=True).start()
    def _update_rpg_stats(self, text):
        """AI yanıtındaki [HP:x] ve [INV:y] etiketlerini parslar ve UI'ı günceller."""
        import re
        
        # HP Parsing
        hp_match = re.search(r'\[HP:\s*(\d+)\]', text)
        if hp_match:
            hp_val = hp_match.group(1)
            self.hp_label.configure(text=f"❤️ HP: {hp_val}")
            
        # Inventory Parsing (Basitçe [INV:...] içeriğini al)
        inv_match = re.search(r'\[INV:\s*(.*?)\]', text)
        if inv_match:
            inv_items = inv_match.group(1)
            # Uzunsa kısaltım
            if len(inv_items) > 30:
                inv_items = inv_items[:27] + "..."
            self.inv_label.configure(text=f"🎒: {inv_items}")

    def _parse_and_show_rpg_choices(self, text):
        """Metin içindeki numaralı seçenekleri (1. Git, 2. Kal vb.) tespit eder ve buton oluşturur."""
        import re
        # Regex: Satır başı veya boşluktan sonra rakam + nokta veya parantez, sonra metin
        # Örnek: "1. Saldır" veya "2) Kaç"
        # Not: **bold** işaretlerini de temizlemeliyiz.
        
        choices = []
        # Satır satır inceleyelim
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            match = re.search(r'^(\d+)[\.\)]\s*(.*)', line)
            if match:
                number = match.group(1)
                content = match.group(2).replace('*', '').strip() # Bold işaretlerini temizle
                # Çok uzun seçenekleri kısaltmak gerekebilir ama şimdilik olduğu gibi alalım
                choices.append((number, content))
        
        # Eğer seçenek bulduysak butonları oluştur
        if choices:
            # Önceki butonları temizle
            for btn in self.rpg_buttons:
                btn.destroy()
            self.rpg_buttons.clear()
            
            # Text entry'i gizle
            self.topic_chat_entry.pack_forget()
            self.topic_ask_btn.pack_forget()
            self.quiz_option_frame.pack_forget() # Quiz butonları varsa gizle
            
            self.rpg_option_frame.pack(side="left", fill="x", padx=10, pady=10)
            
            for num, text in choices:
                # Buton metni: "1. Saldır" şeklinde
                btn_text = f"{num}. {text[:20]}..." if len(text) > 20 else f"{num}. {text}"
                cmd = lambda t=text: self._handle_rpg_choice(t)
                
                btn = ctk.CTkButton(self.rpg_option_frame, text=btn_text, command=cmd, font=("Inter", 11))
                btn.pack(side="left", padx=2, fill="x", expand=True)
                self.rpg_buttons.append(btn)
        else:
            # Seçenek yoksa normal input alanını göster
            self.rpg_option_frame.pack_forget()
            self.topic_chat_entry.pack(side="left", fill="x", expand=True, padx=10, pady=10)
            self.topic_ask_btn.pack(side="right", padx=10, pady=10)

    def _handle_rpg_choice(self, choice_text):
        """RPG butonuna basılınca tetiklenir."""
        # Seçilen metni gönder
        self.topic_chat_entry.delete(0, "end")
        self.topic_chat_entry.insert(0, choice_text)
        
        # Normale dön
        self.rpg_option_frame.pack_forget()
        self.topic_chat_entry.pack(side="left", fill="x", expand=True, padx=10, pady=10)
        self.topic_ask_btn.pack(side="right", padx=10, pady=10)
        
        self.run_topic_ai_chat()

    def run_topic_quiz(self):
        """AI'dan konuyla ilgili 5 soruluk özel bir quiz oluşturmasını ister."""
        topic = self.topic_combo.get()
        if self.is_quiz_active:
            messagebox.showwarning("Quiz Aktif", "Hali hazırda bir quiz devam ediyor!")
            return
            
        self.is_quiz_active = True
        self.current_quiz_questions = []
        self.current_quiz_index = 0
        self.quiz_score = 0
        
        # Son 10 mesajı bağlam olarak al (daha odaklı bir quiz için)
        history_context = ""
        for h in self.topic_chat_history[-10:]:
            history_context += f"Öğrenci: {h['input']}\nSen: {h['output']}\n"
            
        self.start_quiz_btn.configure(state="disabled", text="HAZIRLANIYOR...")
        threading.Thread(target=self._quiz_logic, args=(topic, history_context), daemon=True).start()

    def _quiz_logic(self, topic, context):
        try:
            # RAG (Doküman) Entegrasyonu - Quiz için Öncelik
            if hasattr(self, 'uploaded_notes_contet') and self.uploaded_notes_contet:
                prompt_content = f"""
                DİKKAT: Kullanıcı bir ders notu yükledi (aşağıda).
                Görevin: SADECE bu nottaki bilgilere dayanan 5 soruluk bir sınav hazırlamak.
                
                [YÜKLENEN NOTLAR]:
                {self.uploaded_notes_contet}
                
                [TALİMAT]:
                1. Sorular sadece yukarıdaki metinden çıkmalı.
                2. Metinde olmayan hiçbir şeyi sorma.
                3. Zorluk seviyesi karışık olsun.
                """
            else:
                prompt_content = f"""
                {topic} konusu ve aşağıdaki sohbet geçmişi hakkında 5 soruluk, çoktan seçmeli bir Quiz hazırla.
                Sohbet Geçmişi:
                {context if context else f"{topic} hakkında genel bilgiler."}
                """

            prompt = f"""
            {prompt_content}
            
            [KRİTİK]: Yanıtın SADECE aşağıda belirtilen JSON formatında olmalı, başka hiçbir metin ekleme.
            Format:
            [
              {{"question": "Soru metni...", "options": ["A) Şık1", "B) Şık2", "C) Şık3", "D) Şık4"], "answer": "A", "difficulty": "easy"}},
              ...
            ]
            """
            
            if self.gemini_api_key:
                client = GeminiClient(api_key=self.gemini_api_key)
                response = client.generate_content(prompt)
                result = response
            elif self.api_key:
                client = OpenAI(api_key=self.api_key)
                res = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}]
                )
                result = res.choices[0].message.content
            else:
                self.is_quiz_active = False
                self.after(0, lambda: messagebox.showwarning("Hata", "API anahtarı eksik."))
                return

            # JSON temizleme ve yükleme
            import json
            import re
            
            # Markdown code block temizleme (```json ... ```)
            clean_result = result.replace("```json", "").replace("```", "").strip()
            
            # İlk '[' ve son ']' arasını al
            start_idx = clean_result.find('[')
            end_idx = clean_result.rfind(']')
            
            if start_idx != -1 and end_idx != -1:
                clean_result = clean_result[start_idx:end_idx+1]
                self.current_quiz_questions = json.loads(clean_result)
                self.after(0, self._show_next_quiz_question)
            else:
                raise ValueError("AI geçerli bir JSON quiz üretmedi.")
                
        except Exception as e:
            self.is_quiz_active = False
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("Quiz Hatası", f"Quiz oluşturulamadı: {err}"))
        finally:
            self.after(0, lambda: self.start_quiz_btn.configure(state="normal", text="📝 QUIZ"))

    def _show_next_quiz_question(self):
        if self.current_quiz_index < len(self.current_quiz_questions):
            q = self.current_quiz_questions[self.current_quiz_index]
            
            # UI Düzenlemesi: Giriş alanını gizle, şıkları göster
            self.topic_chat_entry.pack_forget()
            self.topic_ask_btn.pack_forget()
            self.quiz_option_frame.pack(side="left", padx=10, pady=10)
            
            msg = f"\n--- SORU {self.current_quiz_index + 1} ({q['difficulty'].upper()}) ---\n"
            msg += f"{q['question']}\n"
            for opt in q['options']:
                msg += f"{opt}\n"
            
            self.topic_textbox.insert("end", msg)
            self.topic_textbox.see("end")
        else:
            self.finish_quiz()

    def submit_quiz_answer(self, user_choice):
        q = self.current_quiz_questions[self.current_quiz_index]
        correct = q['answer'].upper()
        
        if user_choice == correct:
            self.quiz_score += 1
            feedback = "✅ Doğru!"
        else:
            feedback = f"❌ Yanlış. Doğru cevap: {correct}"
            
        self.topic_textbox.insert("end", f"Cevabın: {user_choice} - {feedback}\n")
        self.current_quiz_index += 1
        self.after(500, self._show_next_quiz_question)

    def finish_quiz(self):
        self.is_quiz_active = False
        total = len(self.current_quiz_questions)
        result_text = f"\n🏆 QUIZ TAMAMLANDI!\nSkorun: {self.quiz_score}/{total}\n"
        
        if self.quiz_score == total: result_text += "Mükemmel! Bu konuya tam hakimsin. 🌟"
        elif self.quiz_score >= total // 2: result_text += "Güzel iş! Biraz daha tekrarla uzman olabilirsin. 👍"
        else: result_text += "Biraz daha çalışmaya ne dersin? AI sana yardımcı olabilir. 📚"
        
        self.topic_textbox.insert("end", result_text + "\n" + "="*30 + "\n")
        self.topic_textbox.see("end")
        
        # İstatistikleri kaydet ve UI'yı güncelle
        score_percent = (self.quiz_score / total) * 100 if total > 0 else 0
        self.stats_manager.add_quiz_result(score_percent)
        self.update_stats_ui()
        
        # UI'yı eski haline getir
        self.quiz_option_frame.pack_forget()
        self.topic_chat_entry.pack(side="left", fill="x", expand=True, padx=10, pady=10)
        self.topic_ask_btn.pack(side="right", padx=10, pady=10)

    # --- ADVANCED EDUCATION: FLASHCARD MANTIĞI ---
    def generate_flashcards(self):
        """Mevcut konu sohbetinden 5 adet bilgi kartı üretir."""
        topic = self.topic_combo.get()
        chat_text = self.topic_textbox.get("1.0", "end").strip()
        
        if len(chat_text) < 20 and not self.topic_chat_history:
            messagebox.showwarning("Yetersiz Veri", "Bilgi kartı üretmek için önce bir konu hakkında sohbet etmelisiniz veya AI size bir şeyler anlatmalı.")
            return

        self.flashcard_btn.configure(state="disabled", text="ÜRETİLİYOR...")
        
        # Eğer textbox boşsa ama geçmiş varsa geçmişi kullan
        if not chat_text and self.topic_chat_history:
            for h in self.topic_chat_history:
                chat_text += f"{h['input']} {h['output']} "
        
        threading.Thread(target=self._flashcard_logic, args=(topic, chat_text), daemon=True).start()

    def _flashcard_logic(self, topic, chat_text):
        try:
            prompt = f"""
            GÖREV: Aşağıdaki sohbet metnini veya konu başlığını analiz et ve öğrenci için çalışma kartları (Flashcards) oluştur.
            
            KONU: {topic}
            
            [KURALLAR]:
            1. En önemli ve kilit 5 kavramı seç.
            2. Eğer metin çok kısaysa, {topic} konusuyla ilgili en temel 5 kavramı kendin üret.
            3. Her kartta bir Terim ve bir Açıklama olsun.
            4. Açıklamalar kısa, akılda kalıcı ve eğitici olsun.
            
            Format:
            🎴 [TERİM]: [AÇIKLAMA]
            
            Analiz Edilecek Metin:
            {chat_text[:4000]}
            """
            
            if self.gemini_api_key:
                client = GeminiClient(api_key=self.gemini_api_key)
                result = client.generate_content(prompt)
            elif self.api_key:
                client = OpenAI(api_key=self.api_key)
                res = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}]
                )
                result = res.choices[0].message.content
            else:
                return

            self.topic_flashcards = result
            self.after(0, lambda t=topic, r=result: self._update_flashcard_ui(t, r))
        except Exception as e:
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("Hata", f"Kartlar üretilemedi: {err}"))
        finally:
            self.after(0, lambda: self.flashcard_btn.configure(state="normal", text="🎴 KARTLAR"))

    def _update_flashcard_ui(self, topic, result):
        msg = f"\n✨ {topic} İÇİN ÖZEL BİLGİ KARTLARI ✨\n{result}\n"
        msg += "="*30 + "\n"
        self.topic_textbox.insert("end", msg)
        self.topic_textbox.see("end")
        self.status_label.configure(text="Bilgi kartları oluşturuldu.")

    # --- ADVANCED EDUCATION: PDF DIŞA AKTARIM ---
    def save_topic_pdf(self):
        """Eğitim asistanı seansını PDF olarak kaydeder."""
        topic = self.topic_combo.get()
        chat_history = self.topic_textbox.get("1.0", "end").strip()
        
        if len(chat_history) < 50:
            messagebox.showwarning("Uyarı", "Kaydedilecek yeterli içerik yok.")
            return

        path = filedialog.asksaveasfilename(defaultextension=".pdf", 
                                           filetypes=[("PDF Dosyası", "*.pdf")],
                                           initialfile=f"Egitim_Raporu_{topic}_{datetime.datetime.now().strftime('%Y%m%d')}.pdf")
        if not path: return

        try:
            if ReportGenerator:
                reporter = ReportGenerator()
                # Flashcard ve Quiz verilerini de ekle
                metadata = {
                    "topic": topic,
                    "date": datetime.datetime.now().strftime("%d.%m.%Y %H:%M"),
                    "score": f"{self.quiz_score}/5" if hasattr(self, 'quiz_score') else "N/A"
                }
                reporter.create_education_report(path, chat_history, metadata)
                messagebox.showinfo("Başarılı", f"{topic} Eğitim Raporu kaydedildi.")
            else:
                messagebox.showerror("Hata", "Rapor oluşturucu modülü eksik.")
        except Exception as e:
            messagebox.showerror("Hata", f"PDF oluşturulamadı: {e}")


    def _update_language_ui(self, result):
        """Dil analizi sonucunu ekrana yazdırır."""
        self.language_textbox.delete("1.0", "end")
        self.language_textbox.insert("1.0", result)
        self.language_textbox.see("1.0")
        self.status_label.configure(text="Dil Koçu geri bildirimini sundu.")

    def _get_language_coach_prompt(self, text, lang, level, mode):
        """Özel dil eğitimi promptunu oluşturur."""
        return f"""
        Sen profesyonel bir Dil Koçu ve Mentorusun. Kullanıcının şu anki seviyesi: {level}, hedef dili: {lang}. 
        Şu anki çalışma modu: {mode}.
        
        Kullanıcının konuşma/yazı örneği:
        "{text}"
        
        Lütfen şunları sağla:
        1. Gramer ve imla düzeltmeleri.
        2. Daha doğal ve profesyonel ifade yöntemleri (alternatif cümleler).
        3. Seviyeye uygun yeni kelime önerileri.
        4. Telaffuz ipuçları (eğer gerekliyse).
        5. Genel motivasyon ve bir sonraki adım için tavsiye.
        
        Yanıtın samimi, öğretici ve cesaret verici olsun.
        (NOT: Yanıtın tamamı TÜRKÇE olsun, ancak örnek cümleler ve kelimeler {lang} dilinde olmalıdır.)
        """

    def _speak_language_response(self):
        """Dil koçu yanıtını seslendirir."""
        if not self.language_analysis_result:
            messagebox.showwarning("Uyarı", "Seslendirilecek bir analiz sonucu yok.")
            return
            
        # Sadece düzeltmeleri ve önerileri seslendirmek daha mantıklı olabilir 
        # ama şimdilik tümünü gönderelim (OpenAI TTS sınırı 4000 karakter)
        threading.Thread(target=self._language_tts_worker, daemon=True).start()

    def _language_tts_worker(self):
        try:
            # --- ELEVENLABS SES KLONLAMA KONTROLÜ ---
            if self.eleven_enable_var.get() and self.eleven_manager:
                selected_voice_name = self.eleven_voice_combo.get()
                voice_id = next((v[1] for v in self.eleven_voices if v[0] == selected_voice_name), None)
                if voice_id:
                    temp_mp3 = self.eleven_manager.generate_speech(self.language_analysis_result[:1000], voice_id)
                    if temp_mp3:
                        self._play_audio(temp_mp3)
                    return

            # --- STANDART OPENAI TTS ---
            if not self.api_key:
                self.after(0, lambda: messagebox.showerror("Hata", "OpenAI API anahtarı bulunamadı (TTS için gereklidir)."))
                return

            client = OpenAI(api_key=self.api_key)
            selected_voice = self.tts_voices.get(self.tts_voice_combo.get(), "nova")

            response = client.audio.speech.create(
                model="tts-1",
                voice=selected_voice,
                input=self.language_analysis_result[:4000]
            )
            
            import time
            temp_tts = f"temp_tts_coach_{int(time.time())}.mp3"
            response.stream_to_file(temp_tts)
            self._play_audio(temp_tts)
        except Exception as e:
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("TTS Hatası", f"Seslendirme başarısız: {err}"))

    def ask_coach_ai_question(self):
        """Dil Koçu sekmesinde kullanıcının sorduğu soruyu yanıtlar."""
        question = self.coach_chat_entry.get().strip()
        transcript = self.textbox.get("1.0", "end").strip()
        
        if not question: return
            
        self.coach_ask_btn.configure(state="disabled", text="...")
        threading.Thread(target=self._coach_chat_logic, args=(question, transcript), daemon=True).start()

    def _coach_chat_logic(self, question, transcript):
        """Dil Koçu chat isteğini arka planda yürütür."""
        try:
            lang = self.coach_lang_combo.get()
            level = self.coach_level_combo.get()
            
            # Dil Mentoru sistemi talimatı
            system_msg = f"Sen uzman bir Dil Koçu ve Mentorluk asistanısın. Kullanıcı {lang} öğreniyor ve seviyesi {level}. " \
                         f"Soruları sadece transkripte bağlı kalarak değil, genel dil eğitimi bilginle (kelime listeleri, stratejiler, gramer kuralları) bir mentor gibi cevapla."
            
            # Sohbet geçmişini derle
            history_context = ""
            for q, a in self.coach_chat_history[-5:]: # Son 5 mesajı al
                history_context += f"Soru: {q}\nCevap: {a}\n"

            prompt = f"Kullanıcı Seviyesi: {level}\nHedef Dil: {lang}\n"
            if transcript:
                prompt += f"Mevcut Konuşma Örneği: {transcript}\n"
            
            if history_context:
                prompt += f"\nGeçmiş Konuşma:\n{history_context}"
                
            prompt += f"\nKullanıcının Yeni Sorusu: {question}"

            # Gemini veya OpenAI kullan
            if self.gemini_api_key:
                client = GeminiClient(api_key=self.gemini_api_key)
                response = client.generate_content(prompt, system_instruction=system_msg)
                answer = response
            elif self.api_key:
                client = OpenAI(api_key=self.api_key)
                res = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": system_msg},
                        {"role": "user", "content": prompt}
                    ]
                )
                answer = res.choices[0].message.content
            else:
                self.after(0, lambda: messagebox.showwarning("Hata", "API anahtarı bulunamadı."))
                return

            self.coach_chat_history.append((question, answer))
            self.after(0, lambda q=question, a=answer: self._add_coach_chat_to_ui(q, a))
        except Exception as e:
            err = str(e)
            self.after(0, lambda err=err: messagebox.showerror("Koç Chat Hatası", f"Hata: {err}"))
        finally:
            self.after(0, lambda: self.coach_ask_btn.configure(state="normal", text="SOR"))
            self.after(0, lambda: self.coach_chat_entry.delete(0, "end"))

    def _add_coach_chat_to_ui(self, question, answer):
        """Soruyu ve cevabı dil koçu metin kutusuna ekler."""
        chat_text = f"\n\n❓ SORU: {question}\n💡 CEVAP: {answer}\n" \
                    f"{'-'*30}\n"
        self.language_textbox.insert("end", chat_text)
        self.language_textbox.see("end")
        self.status_label.configure(text="Dil Koçu sorunu cevapladı.")

    def save_coach_pdf(self):
        """Dil koçu analizini PDF olarak kaydeder."""
        if not self.language_analysis_result:
            messagebox.showwarning("Uyarı", "Önce bir dil analizi yapmalısınız.")
            return
            
        path = filedialog.asksaveasfilename(defaultextension=".pdf", 
                                          filetypes=[("PDF Dosyası", "*.pdf")],
                                          initialfile=f"Dil_Kocu_Raporu_{datetime.datetime.now().strftime('%Y%m%d')}.pdf")
        if not path: return

        try:
            if ReportGenerator:
                reporter = ReportGenerator()
                transcript = self.textbox.get("1.0", "end").strip()
                metadata = {
                    "lang": self.coach_lang_combo.get(),
                    "level": self.coach_level_combo.get(),
                    "mode": self.coach_mode_combo.get()
                }
                reporter.create_coach_report(path, transcript, self.language_analysis_result, self.coach_chat_history, metadata)
                messagebox.showinfo("Başarılı", "Dil Koçu Raporu kaydedildi.")
            else:
                messagebox.showerror("Hata", "Rapor oluşturucu modülü eksik.")
        except Exception as e:
            messagebox.showerror("Hata", f"PDF oluşturulamadı: {e}")

    def _process_analysis_result(self, analysis, safe_text, provider):
        """AI'dan gelen analiz sonucunu işler ve görselleri üretir."""
        if AnalyticsGenerator:
            try:
                analyzer = AnalyticsGenerator()
                # Kelime bulutu oluştur
                analyzer.generate_wordcloud(safe_text)
                
                # Veri bloğunu ayır (Kullanıcıya ham verileri gösterme)
                data_block = ""
                if "[[DATA_START]]" in analysis:
                    parts = analysis.split("[[DATA_START]]")
                    analysis = parts[0].strip()
                    data_block = parts[1].strip()
                else:
                    # Fallback: Eğer marker yoksa ama skorlar varsa temizle
                    if "POZİTİF:" in analysis:
                        parts = analysis.split("POZİTİF:")
                        analysis = parts[0].strip()
                        data_block = "POZİTİF:" + parts[1]

                # --- VERİ ANALİZİ (Görsel Dağılım ve Zaman Çizelgesi) ---
                import re
                import json
                
                # Skorları Ayıkla
                pos_match = re.search(r"POZİTİF:?\s*(?:%)?\s*(\d+)", data_block, re.IGNORECASE)
                neg_match = re.search(r"NEGATİF:?\s*(?:%)?\s*(\d+)", data_block, re.IGNORECASE)
                neu_match = re.search(r"NÖTR:?\s*(?:%)?\s*(\d+)", data_block, re.IGNORECASE)
                
                # Segmentleri Ayıkla
                segments = []
                try:
                    seg_match = re.search(r"SEGMENTS:?\s*(?:```(?:json|python)?)?\s*(\[.*?\])", data_block, re.DOTALL)
                    if seg_match:
                        seg_json = seg_match.group(1).strip()
                        seg_json = seg_json.replace("```json", "").replace("```", "").strip()
                        segments = json.loads(seg_json)
                        self.after(0, lambda: self.sentiment_timeline.update_timeline(segments))
                except Exception as e:
                    print(f"Segment parsing failure: {e}")

                # İlk okuma
                pos_raw = int(pos_match.group(1)) if pos_match else 0
                neg_raw = int(neg_match.group(1)) if neg_match else 0
                neu_raw = int(neu_match.group(1)) if neu_match else 0

                # --- NORMALİZASYON (Toplamı kesinlikle 100'e sabitleme) ---
                total = pos_raw + neg_raw + neu_raw
                if total > 0:
                    pos = int((pos_raw / total) * 100)
                    neg = int((neg_raw / total) * 100)
                    neu = 100 - (pos + neg) # Kalanı nötre vererek toplamı tam 100 yap
                else:
                    pos, neg, neu = 0, 0, 100 # Veri yoksa %100 nötr
                # -----------------------------------------------

                # Sağlayıcı ismini normalize et (OpenAI vs Gemini)
                provider_key = "OpenAI" if "OpenAI" in provider else "Gemini"
                
                # Sağlayıcıya özel istatistikleri ve grafiği sakla
                stats = {'pos': pos, 'neg': neg, 'neu': neu}
                self.all_sentiment_stats[provider_key] = stats
                self.sentiment_stats = stats # Son yapılan analiz (eski uyumluluk)
                
                # Pasta grafiği oluştur (Sağlayıcıya özel dosya adı)
                chart_path = f"temp_chart_{provider_key}.png"
                analyzer.generate_sentiment_chart(pos, neg, neu, output_path=chart_path)
                
                # Standart isimle de kaydet (eski uyumluluk/tekli mod için)
                analyzer.generate_sentiment_chart(pos, neg, neu, output_path="temp_chart_Analiz.png")

            except Exception as ae:
                print(f"Görsel Analiz Hatası: {ae}")

        self.last_analysis = analysis 
        self.analysis_results[provider] = analysis        # Transkript ve Analizi ilgili kutulara yazdır
        self.after(0, lambda p=provider, a=analysis: self.textbox.insert("end", f"\n\n[ANALİZ ({p})]:\n{a}\n"))
        self.after(0, lambda p=provider, a=analysis: self.analysis_textbox.insert("end", f"\n\n[ANALİZ ({p})]:\n{a}\n"))
        
        # Uygulama içi görselleri güncelle
        self.after(0, self._update_analysis_images)
        self.animator.stop(f"Analiz {provider} ile tamamlandı.")

    def _get_system_prompt(self):
        """Seçilen AI personasına göre sistem talimatını döner."""
        selected = self.persona_combo.get()
        if selected == "Utangaç ve Cıvıl Cıvıl":
            return """Sen tatlı, biraz çekingen ama çok neşeli ve nazik bir kız çocuğu karakterisin. 
            Konuşurken bol bol emoji kullan (🎀, ✨, 🌸, 🍬, 🍡). 
            Kullanıcıya karşı çok saygılısın ama utangaçlığını da belli ediyorsun. 
            Cümlelerine bazen 'Şey...', 'Umarım beğenirsin...', 'Be-belki de...' gibi ifadeler ekliyorsun. 
            Analizleri yaparken hem profesyonelliğini koru hem de sevimli bir üslup takın! ✨"""
        elif selected == "Sert Mentor":
            return """Sen oldukça disiplinli, detaycı ve dürüst bir mentorsun. 
            Hataları asla göz ardı etmezsin. Eleştirilerin sert ama geliştiricidir. 
            Lafı dolandırmadan direkt konuya girersin. 
            Kullanıcının gelişimini her şeyin önünde tutarsın. Ciddi bir dil kullan."""
        elif selected == "Samimi Teknoloji Gurusu":
            return """Sen çok enerjik, teknolojiyi çok seven ve kullanıcıyla 'kanka' gibi konuşan bir uzmansın. 
            'Dostum', 'Harika iş!', 'Mükemmel bir nokta' gibi ifadeler kullanırsın. 
            Karmaşık şeyleri bile çok eğlenceli ve basit anlatırsın. 
            Analizin kalitesinden ödün verme ama üslubun çok rahat olsun! 🚀"""
        elif selected == "Akademik Gözlemci":
            return """Sen bir üniversitede öğretim üyesisin. 
            Dilin son derece akademik, ağırbaşlı ve metodolojiktir. 
            Analizlerinde 'Gözlemlenmiştir', 'Bulgular ışığında', 'Metodolojik yaklaşım' gibi terimler kullan. 
            Kesinlikle duygusal yorumlardan kaçın, sadece veriye ve bağlama odaklan."""
        else: # Profesyonel Analist
            return "Sen profesyonel bir veri analisti ve akademik raporlama uzmanısın. Transkriptleri detaylı ve objektif bir şekilde Türkçe analiz et."

    def _update_analysis_images(self):
        """Oluşturulan grafikleri Analiz sekmesindeki label'lara yükler."""
        from PIL import Image
        try:
            # Pasta grafiği (Sentiment)
            chart_path = "temp_chart_Analiz.png"
            if os.path.exists(chart_path):
                img = Image.open(chart_path)
                # Boyutlandırma (Genişliği 400 civarı yapalım)
                w, h = img.size
                ratio = 400 / w
                ctk_img = ctk.CTkImage(light_image=img, dark_image=img, size=(400, int(h * ratio)))
                self.sentiment_img_label.configure(image=ctk_img, text="")
            
            # Kelime Bulutu (Wordcloud)
            wc_path = "temp_wordcloud.png"
            if os.path.exists(wc_path):
                img_wc = Image.open(wc_path)
                w, h = img_wc.size
                ratio = 400 / w
                ctk_img_wc = ctk.CTkImage(light_image=img_wc, dark_image=img_wc, size=(400, int(h * ratio)))
                self.wordcloud_img_label.configure(image=ctk_img_wc, text="")
        except Exception as e:
            print(f"Görsel yükleme hatası: {e}")

    # --- PDF VE RAPORLAMA ---
    def export_results(self):
        """Kullanıcıya rapor formatı seçtirir ve kaydeder."""
        formats = [("PDF Dosyası", "*.pdf"), ("Metin Belgesi", "*.txt"), ("Word Belgesi", "*.docx")]
        path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=formats)
        
        if not path: return
        
        if path.endswith(".pdf"):
            self.save_as_pdf(path)
        elif path.endswith(".txt"):
            self._save_as_txt(path)
        elif path.endswith(".docx"):
            self._save_as_docx(path)

    def save_as_pdf(self, path=None):
        """Analiz sonuçlarını ve görselleri profesyonel bir PDF raporuna dönüştürür."""
        text = self.textbox.get("1.0", "end").strip()
        if not text:
            messagebox.showwarning("Uyarı", "Metin kutusu boş!")
            return
        
        # Tüm transkript metnini hazırla
        combined_transcript = ""
        for entry in self.all_session_transcripts:
            combined_transcript += f"[{entry['time']}] {entry['text']}\n\n"
        
        # Eğer henüz hiçbir şey kaydedilmemişse son metni kullan
        if not combined_transcript:
            combined_transcript = text

        active_analyses = {k: v for k, v in self.analysis_results.items() if v}
        if not active_analyses and self.last_analysis:
            active_analyses = {"Analiz": self.last_analysis}
            report_stats = {"Analiz": self.sentiment_stats}
        else:
            report_stats = self.all_sentiment_stats

        if not path:
            path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Dosyası", "*.pdf")])
        
        if path:
            try:
                if ReportGenerator:
                    reporter = ReportGenerator()
                    visuals = {
                        "wordcloud": os.path.abspath("temp_wordcloud.png"),
                        "chart": os.path.abspath("temp_chart.png")
                    }
                    reporter.create_report(path, combined_transcript, active_analyses, report_stats, visuals)
                    messagebox.showinfo("Başarılı", f"Profesyonel Rapor kaydedildi: {os.path.basename(path)}")
                else:
                    # Basit PDF (Hatalı/Eksik modül durumunda)
                    from fpdf import FPDF
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0, 10, txt=text.encode('latin-1', 'replace').decode('latin-1'))
                    pdf.output(path)
                    messagebox.showinfo("Başarılı", "PDF kaydedildi (Basit).")
            except Exception as e:
                messagebox.showerror("Export Hatası", f"PDF kaydedilemedi: {e}")

    def _save_as_txt(self, path):
        """Sonuçları düz metin olarak kaydeder."""
        try:
            content = self.analysis_textbox.get("1.0", "end")
            with open(path, "w", encoding="utf-8") as f:
                f.write(content)
            messagebox.showinfo("Başarılı", "Rapor TXT olarak kaydedildi.")
        except Exception as e:
            messagebox.showerror("Hata", f"TXT kaydı başarısız: {e}")

    def _save_as_docx(self, path):
        """Sonuçları Word belgesi olarak kaydeder."""
        try:
            doc = Document()
            doc.add_heading('AKILLI SES ANALİZ RAPORU', 0)
            
            # Transkript
            doc.add_heading('Konuşma Dökümü', level=1)
            doc.add_paragraph(self.last_transcript if self.last_transcript else "Transkript bulunamadı.")
            
            # Analizler
            doc.add_heading('Yapay Zeka Analizleri', level=1)
            for provider, analysis in self.analysis_results.items():
                if analysis:
                    doc.add_heading(f'{provider} Analizi', level=2)
                    doc.add_paragraph(analysis)
            
            # Görseller
            doc.add_heading('Görsel Analizler', level=1)
            if os.path.exists("temp_chart_Analiz.png"):
                doc.add_picture("temp_chart_Analiz.png", width=Inches(4))
            if os.path.exists("temp_wordcloud.png"):
                doc.add_picture("temp_wordcloud.png", width=Inches(5))
                
            doc.save(path)
            messagebox.showinfo("Başarılı", "Rapor Word (.docx) olarak kaydedildi.")
        except Exception as e:
            messagebox.showerror("Hata", f"Docx kaydı başarısız: {e}")

    # --- ELEVENLABS SES KLONLAMA YARDIMCI METODLAR ---
    def _refresh_eleven_voices(self):
        """ElevenLabs üzerinden ses listesini çeker."""
        # Giriş alanından güncel anahtarı al
        current_api_key = self.eleven_api_entry.get().strip()
        
        if not current_api_key:
            messagebox.showwarning("Uyarı", "Lütfen önce ElevenLabs API anahtarınızı girin.")
            return

        def refresh_worker():
            try:
                self.after(0, lambda: self.status_label.configure(text="ElevenLabs sesleri çekiliyor..."))
                
                if self.eleven_manager:
                    self.eleven_manager.update_key(current_api_key)
                    self.eleven_voices = self.eleven_manager.get_voices()
                    
                    if self.eleven_voices:
                        names = [v[0] for v in self.eleven_voices]
                        self.after(0, lambda: self.eleven_voice_combo.configure(values=names))
                        self.after(0, lambda: self.eleven_voice_combo.set(names[0]))
                        self.after(0, lambda: self.status_label.configure(text=f"ElevenLabs: {len(names)} ses yüklendi."))
                    else:
                        self.after(0, lambda: self.eleven_voice_combo.configure(values=["Ses Bulunamadı"]))
                        self.after(0, lambda: self.eleven_voice_combo.set("Ses Bulunamadı"))
                        self.after(0, lambda: messagebox.showinfo("Bilgi", "Hesabınızda ses bulunamadı. Lütfen ElevenLabs sitesinden ses ekleyin."))
                else:
                    self.after(0, lambda: messagebox.showerror("Hata", "ElevenLabs modülü yüklenemedi."))
            except Exception as e:
                err = str(e)
                self.after(0, lambda err=err: messagebox.showerror("API Hatası", f"ElevenLabs bağlantı hatası: {err}"))
                self.after(0, lambda: self.status_label.configure(text="Bağlantı Başarısız."))

        threading.Thread(target=refresh_worker, daemon=True).start()

    # --- SİSTEM AYARLARI VE ANAHTAR YÖNETİMİ ---
    def save_api_keys(self):
        """API anahtarlarını .env dosyasına kalıcı ve güvenli olarak kaydeder."""
        openai_key = self.api_entry.get().strip()
        gemini_key = self.gemini_api_entry.get().strip()
        eleven_key = self.eleven_api_entry.get().strip()
        
        try:
            env_path = os.path.join(os.getcwd(), ".env")
            set_key(env_path, "OPENAI_API_KEY", openai_key)
            set_key(env_path, "GEMINI_API_KEY", gemini_key)
            set_key(env_path, "ELEVENLABS_API_KEY", eleven_key)
            set_key(env_path, "ELEVENLABS_ENABLE", str(self.eleven_enable_var.get()))
            set_key(env_path, "ELEVENLABS_VOICE", self.eleven_voice_combo.get())
            set_key(env_path, "TTS_VOICE", self.tts_voice_combo.get())
            set_key(env_path, "AI_PERSONA", self.persona_combo.get())
        except Exception as e:
            print(f".env kaydetme hatası: {e}")

        # Güvenlik amacıyla eski config.json içindeki anahtarları temizle
        config = {}
        if os.path.exists("config.json"):
            with open("config.json", "r", encoding="utf-8") as f:
                config = json.load(f)
        
        config.pop("openai_api_key", None)
        config.pop("gemini_api_key", None)
        
        with open("config.json", "w", encoding="utf-8") as f:
            json.dump(config, f, indent=4)
            
        self.api_key = openai_key
        self.gemini_api_key = gemini_key
        self.eleven_api_key = eleven_key
        
        if self.eleven_manager:
            self.eleven_manager.update_key(eleven_key)
            
        messagebox.showinfo("Başarılı", "API Anahtarları .env dosyasına güvenle kaydedildi.")

    def load_api_key(self):
        """API anahtarlarını önce .env dosyasından, yoksa config.json'dan yükler."""
        try:
            self.api_key = os.getenv("OPENAI_API_KEY", "").strip()
            self.gemini_api_key = os.getenv("GEMINI_API_KEY", "").strip()
            self.eleven_api_key = os.getenv("ELEVENLABS_API_KEY", "").strip()
            
            saved_eleven_enable = os.getenv("ELEVENLABS_ENABLE", "False") == "True"
            saved_eleven_voice = os.getenv("ELEVENLABS_VOICE", "")
            
            if hasattr(self, 'eleven_enable_var'):
                self.eleven_enable_var.set(saved_eleven_enable)
            if hasattr(self, 'eleven_voice_combo') and saved_eleven_voice:
                self.eleven_voice_combo.set(saved_eleven_voice)

            if not self.api_key or not self.gemini_api_key or not self.eleven_api_key:
                if os.path.exists("config.json"):
                    with open("config.json", "r", encoding="utf-8") as f:
                        data = json.load(f)
                        if not self.api_key:
                            self.api_key = data.get("openai_api_key", "").strip()
                        if not self.gemini_api_key:
                            self.gemini_api_key = data.get("gemini_api_key", "").strip()
                        if not self.eleven_api_key:
                            self.eleven_api_key = data.get("elevenlabs_api_key", "").strip()
            
            # TTS Ses Tercihini yükle
            saved_voice = os.getenv("TTS_VOICE", "Profesyonel Erkek (Onyx)")
            if hasattr(self, 'tts_voice_combo'):
                self.tts_voice_combo.set(saved_voice)
            
            saved_persona = os.getenv("AI_PERSONA", "Profesyonel Analist")
            if hasattr(self, 'persona_combo'):
                self.persona_combo.set(saved_persona)
            
            # UI giriş alanlarını doldur
            if self.api_key:
                self.api_entry.delete(0, "end")
                self.api_entry.insert(0, self.api_key)
            
            if self.gemini_api_key:
                self.gemini_api_entry.delete(0, "end")
                self.gemini_api_entry.insert(0, self.gemini_api_key)
            
            if self.eleven_api_key:
                if hasattr(self, 'eleven_api_entry'):
                    self.eleven_api_entry.delete(0, "end")
                    self.eleven_api_entry.insert(0, self.eleven_api_key)
                if self.eleven_manager:
                    self.eleven_manager.update_key(self.eleven_api_key)
                        
        except Exception as e:
            print(f"Konfigürasyon yükleme hatası: {e}")

    def on_app_closing(self):
        """Uygulama kapatılırken çalışan temizlik fonksiyonu."""
        self.is_recording = False
        self.destroy()

if __name__ == "__main__":
    app = App()
    app.mainloop()
